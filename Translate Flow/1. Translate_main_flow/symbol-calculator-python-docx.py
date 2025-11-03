# -*- coding: utf-8 -*-
import os
import time
import traceback
from dotenv import load_dotenv

# import unicodedata # Больше не нужен
# import tempfile # Не нужен

# --- Библиотеки ---
try:
    import docx

    try:
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError:
        PackageNotFoundError = None
except ImportError:
    print("ОШИБКА: python-docx не найден. Установите: pip install python-docx");
    exit()

try:
    import aspose.words as aw
except ImportError:
    print("ОШИБКА: aspose-words не найден. Установите: pip install aspose-words");
    exit()

try:
    import pandas as pd
except ImportError:
    print("ОШИБКА: pandas не найден. Установите: pip install pandas openpyxl");
    exit()  # openpyxl нужен для записи в .xlsx


# --- Функция count_total_files (без изменений) ---
def count_total_files(folder_path):
    total_files = 0
    for root, _, files in os.walk(folder_path):
        for filename in files:
            if filename.lower().endswith(('.doc', '.docx')) and not filename.startswith('~$'):
                total_files += 1
    return total_files


# --- Функция подсчета python-docx (без изменений) ---
def count_chars_python_docx(filepath: str) -> int:
    """Подсчет символов для .docx с помощью python-docx. Возвращает >= 0 или вызывает исключение."""
    char_count = 0
    try:
        document = docx.Document(filepath)
        for para in document.paragraphs: char_count += len(para.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs: char_count += len(para.text)
        return char_count  # Возвращаем 0 или больше
    except (PackageNotFoundError if PackageNotFoundError else Exception) as e:
        print(f"\nОшибка python-docx при чтении {os.path.basename(filepath)}: {e}", flush=True)
        raise RuntimeError(f"python-docx failed: {e}") from e
    except Exception as e:
        print(f"\nНеожиданная ошибка python-docx для {os.path.basename(filepath)}: {e}", flush=True)
        raise RuntimeError(f"python-docx unexpected error: {e}") from e


# --- Функция подсчета Aspose range.text (без изменений) ---
def count_chars_aspose_range_text(doc: aw.Document) -> int:
    """Считает символы через doc.range.text, сохраняя пробелы, \n, \r, \t."""
    try:
        text = doc.range.text
        cleaned_text = text.replace(chr(7), "").replace(chr(11), "").replace(chr(12), "")
        char_count = len(cleaned_text)
        print(f"DEBUG: Посчитано через Range.Text Fallback: {char_count}", flush=True)
        return char_count
    except Exception as e:
        print(f"\nОшибка Aspose range.text: {e}", flush=True)
        # При ошибке range.text возвращаем -1 как маркер неудачи
        # Это позволит отличить неудачу от реального 0
        return -1  # Маркер ошибки


# --- Основная функция подсчета (Новая последовательность) (без изменений) ---
def count_docs_in_folder(folder_path, license_path):
    report_data = [];
    processed_files = 0;
    total_chars_overall = 0
    # --- Лицензия ---
    aspose_license_applied = False
    try:
        if not os.path.exists(license_path):
            print(f"Предупреждение: Лицензия Aspose не найдена: {license_path}.")
        else:
            license = aw.License(); license.set_license(license_path); print(
                "Лицензия Aspose.Words успешно применена."); aspose_license_applied = True
    except Exception as e:
        print(f"Ошибка лицензии Aspose: {e}. Aspose будет работать в режиме оценки.")
    print(
        f"Используется python-docx для .docx; Aspose {'с лицензией' if aspose_license_applied else 'в режиме оценки'} для .doc и fallback.")
    # --- Подсчет файлов ---
    total_files = count_total_files(folder_path)
    if total_files == 0: print("Файлы .doc или .docx не найдены."); return [], 0, 0
    print(f"Всего .doc/.docx файлов для обработки: {total_files}")

    # --- Основной цикл ---
    for root, _, files in os.walk(folder_path):
        files.sort()
        for filename in files:
            if not filename.lower().endswith(('.doc', '.docx')) or filename.startswith('~$'): continue

            filepath = os.path.join(root, filename)
            char_count = 0;
            doc_count = 0;
            threshold_exceeded = 0
            error_message = None;
            method_used = "N/A"
            relative_path = os.path.relpath(filepath, folder_path)
            start_time_file = time.time()
            use_aspose = False
            aspose_doc = None

            # --- Определение первичного метода ---
            if filename.lower().endswith('.docx'):
                try:
                    # 1. Попытка python-docx
                    docx_count = count_chars_python_docx(filepath)
                    if docx_count > 0:
                        char_count = docx_count
                        method_used = "python-docx"
                        print(f"\nDEBUG: {filename} - Использован python-docx: {char_count}", flush=True)
                    else:
                        # python-docx сработал, но дал 0 -> переходим к Aspose
                        print(f"INFO: {filename} - python-docx дал 0. Используем Aspose Fallback...", flush=True)
                        use_aspose = True
                        error_message = "python-docx returned 0"  # Информационное сообщение

                except Exception as e_docx:
                    # python-docx упал -> переходим к Aspose
                    print(f"INFO: {filename} - Ошибка python-docx ({e_docx}). Используем Aspose Fallback...",
                          flush=True)
                    use_aspose = True
                    error_message = f"python-docx failed: {e_docx}"  # Записываем ошибку

            elif filename.lower().endswith('.doc'):
                use_aspose = True  # Для .doc сразу Aspose
            else:
                continue

            # --- Выполнение Aspose (если нужно) ---
            if use_aspose:
                original_error_message = error_message  # Сохраняем ошибку python-docx, если была
                aspose_method_successful = False
                try:
                    aspose_doc = aw.Document(filepath)

                    # --- Попытка 1: Properties ---
                    count_properties = -1  # Маркер "не сработало"
                    try:
                        prop_val = aspose_doc.built_in_document_properties.characters_with_spaces
                        if prop_val > 0:  # Используем только если > 0
                            count_properties = prop_val
                            print(f"DEBUG: {filename} - Aspose Properties: {count_properties}", flush=True)
                    except Exception as prop_ex:
                        print(f"\nПредупреждение: Не удалось прочитать Aspose Properties для {filename}: {prop_ex}",
                              flush=True)
                        # count_properties остается -1

                    if count_properties > 0:
                        char_count = count_properties
                        if filename.lower().endswith('.doc'):
                            method_used = "Aspose (DOC - Properties)"
                        else:
                            method_used = "Aspose (Fallback - Properties)"; error_message = None  # Успех Aspose
                        aspose_method_successful = True
                        print(f"DEBUG: {filename} - Использован результат Aspose Properties: {char_count}", flush=True)

                    # --- Попытка 2: Range.Text (если Properties не сработал) ---
                    if not aspose_method_successful:
                        print(f"INFO: {filename} - Aspose Properties <= 0 или ошибка. Попытка Range.Text Fallback...",
                              flush=True)
                        count_range = count_chars_aspose_range_text(aspose_doc)

                        if count_range >= 0:  # Если range.text не вернул маркер ошибки (-1)
                            char_count = count_range  # Используем результат, даже если 0
                            if filename.lower().endswith('.doc'):
                                method_used = "Aspose (DOC - Range.Text)"
                            else:
                                method_used = "Aspose (Fallback - Range.Text)"; error_message = None  # Успех Aspose
                            aspose_method_successful = True  # Считаем range.text успехом, даже если 0
                            print(f"DEBUG: {filename} - Использован результат Range.Text Fallback: {char_count}",
                                  flush=True)
                            if char_count == 0:
                                print(f"WARNING: {filename} - И Properties, и Range.Text дали 0 или ошибку.")
                        else:  # Ошибка в count_chars_aspose_range_text
                            method_used = "Error"
                            if original_error_message:
                                error_message = original_error_message + f"; Aspose Range.Text Error"
                            else:
                                error_message = "Aspose Range.Text Error"
                            char_count = 0

                except Exception as e_aspose_load:  # Ошибка при ЗАГРУЗКЕ документа Aspose
                    print(f"\nОшибка ЗАГРУЗКИ Aspose для {filename}: {e_aspose_load}", flush=True)
                    if original_error_message:
                        error_message = original_error_message + f"; Aspose Load Error: {e_aspose_load}"
                    else:
                        error_message = f"Aspose Load Error: {e_aspose_load}"
                    method_used = "Error";
                    char_count = 0

            # --- Расчет 'документов' ---
            if char_count < 0: char_count = 0  # Если что-то пошло не так
            total_chars_overall += char_count
            threshold = 50000
            if char_count > 0: threshold_exceeded = (char_count - 1) // threshold; doc_count = threshold_exceeded + 1

            # --- Добавление в отчет ---
            report_data.append(
                [relative_path, char_count, threshold_exceeded, doc_count, method_used, error_message or ""])

            # --- Прогресс ---
            processed_files += 1;
            progress = (processed_files / total_files) * 100
            time_spent_file = time.time() - start_time_file
            print(
                f"\nОбработано: {processed_files}/{total_files} ({progress:.1f}%) | {time_spent_file:.2f}с | Файл: {filename:<40}",
                end='\r', flush=True)

    # --- Завершение ---
    print("\n" + "=" * 80)
    print("Обработка завершена.")
    return report_data, total_files, total_chars_overall


# --- Обновленная функция для записи отчета в Excel ---
def write_report_to_excel(report_data, total_files, total_chars_overall, folder_path):
    # 1. Сначала рассчитываем total_calculated_docs, так как он нужен для имени файла
    total_calculated_docs = 0
    # Подсчет итоговых 'документов' из report_data (список списков)
    for row in report_data:
        try:
            # row[3 Translate_politics] это "Итого 'документов'"
            if len(row) > 3 and row[3] is not None:
                total_calculated_docs += int(row[3])
        except (ValueError, TypeError):
            pass  # Пропускаем, если значение не может быть преобразовано в int

    # 2. Формируем новое имя файла
    # Формат: ДД.ММ.ГГГГ_ИтогоДокументов.xlsx
    current_date_str = time.strftime('%d.%m.%Y')
    report_filename = f"{current_date_str}_{total_calculated_docs}.xlsx"
    excel_filepath = os.path.join(folder_path, report_filename)  # Отчет сохранится в указанной папке

    headers = ["Относительный путь", "Количество символов", "Превышений порога (50k)", "Итого 'документов'",
               "Метод подсчета", "Примечание/Ошибка"]

    # Преобразование данных в DataFrame pandas
    df_report = pd.DataFrame(report_data, columns=headers)

    # Создание DataFrame для итоговой информации
    summary_data = {
        "Описание": [
            "Итого файлов обработано:",
            "Итого символов (суммарно):",
            "Итого 'документов' (расчетное):"  # Это значение уже есть в total_calculated_docs
        ],
        "Значение": [
            total_files,
            total_chars_overall,
            total_calculated_docs  # Используем уже посчитанное значение
        ]
    }
    df_summary = pd.DataFrame(summary_data)

    try:
        with pd.ExcelWriter(excel_filepath, engine='openpyxl') as writer:
            # Записываем основную таблицу отчета
            df_report.to_excel(writer, sheet_name='Отчет', index=False)

            # Определяем, с какой строки начать запись итоговой информации
            start_row_for_summary = len(df_report) + 2

            # Записываем итоговую информацию ниже основной таблицы
            df_summary.to_excel(writer, sheet_name='Отчет', startrow=start_row_for_summary, index=False, header=False)

        print(f"Отчет сохранен в: {excel_filepath}")
    except PermissionError:
        print(f"\nОшибка записи: Недостаточно прав для сохранения отчета в {excel_filepath}.")
    except Exception as e:
        print(f"\nОшибка при записи в Excel файл {excel_filepath}: {e}")
        traceback.print_exc()


# --- Основная функция main (без изменений) ---
def main():
    # Загружаем переменные окружения из .env файла
    load_dotenv(os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '.env'))
    
    folder_path = input("Введите путь к каталогу с документами: ").strip().strip('"')
    if not os.path.isdir(folder_path): print(f"Ошибка: Путь '{folder_path}' не каталог."); return
    
    # Получаем путь к лицензии из переменной окружения
    license_path = os.getenv('ASPOSE_LICENSE_PATH')
    if not license_path:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        license_path = os.path.join(script_dir, "Aspose.Words.lic")
    
    start_time_global = time.time()
    try:
        report_data, total_files, total_chars = count_docs_in_folder(folder_path, license_path)
        if total_files > 0:
            write_report_to_excel(report_data, total_files, total_chars, folder_path)  # Вызов остается тем же
        else:
            print("Не найдено файлов .doc или .docx для обработки.")
    except Exception as e:
        print(f"\nГЛОБАЛЬНАЯ ошибка: {e}");
        traceback.print_exc()
    finally:
        print(f"Общее время выполнения: {time.time() - start_time_global:.2f} секунд.")


# --- Точка входа (без изменений) ---
if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"Критическая ошибка: {e}")
    input("\nНажмите Enter для выхода...")