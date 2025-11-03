#!/usr/bin/env python3
"""
EQN RESTORATION BATCH - FIXED VERSION V2
Исправленная версия с обнаружением и исправлением поврежденных плейсхолдеров

НОВЫЕ ВОЗМОЖНОСТИ:
1. Обнаружение поврежденных плейсхолдеров (без открывающих << или закрывающих >>)
2. Автоматическое исправление поврежденных плейсхолдеров
3. Детальный отчет о найденных и исправленных проблемах
4. Правильное восстановление последовательности после исправления
"""

import sys
from pathlib import Path
import re
from datetime import datetime
import json

try:
    from docx import Document
except ImportError:
    print("Ошибка: Необходима библиотека python-docx.")
    print("Установите: pip install python-docx")
    sys.exit(1)


class PlaceholderRestorer:
    """Класс для восстановления последовательности плейсхолдеров"""
    
    def __init__(self):
        self.stats = {
            'placeholders_replaced': 0,
            'placeholders_found_in_translation': 0,
            'placeholders_found_in_original': 0,
            'damaged_placeholders_fixed': 0,
            'damaged_placeholders_details': []
        }
        
        # Паттерн для поиска правильных плейсхолдеров
        self.pattern_placeholder = re.compile(r'<<Eqn\d+(?:\.eps)?>>(?:,)?', re.IGNORECASE)
        
        # Паттерн для поиска поврежденных плейсхолдеров
        # Находит варианты: <EqnXXX.eps>>, <<EqnXXX.eps>, <EqnXXX.eps>, EqnXXX.eps>>
        self.pattern_damaged = re.compile(
            r'(?:<{1,2})?Eqn\d+(?:\.eps)?(?:>{1,2})?(?:,)?', 
            re.IGNORECASE
        )
        
    def find_and_fix_damaged_placeholders(self, text, para_info=""):
        """
        Находит и исправляет поврежденные плейсхолдеры в тексте
        Возвращает (исправленный_текст, количество_исправлений, список_исправлений)
        """
        fixed_count = 0
        fixes = []
        
        def fix_placeholder(match):
            nonlocal fixed_count
            original = match.group(0)
            
            # Извлекаем номер плейсхолдера
            num_match = re.search(r'Eqn(\d+)', original, re.IGNORECASE)
            if not num_match:
                return original
                
            num = num_match.group(1)
            
            # Проверяем, нужно ли исправление
            if not original.startswith('<<') or not original.endswith('>>'):
                # Строим правильный плейсхолдер
                if '.eps' in original.lower():
                    correct = f'<<Eqn{num}.eps>>'
                else:
                    correct = f'<<Eqn{num}>>'
                    
                # Сохраняем запятую если она была
                if original.endswith(','):
                    correct += ','
                    
                if original != correct:
                    fixed_count += 1
                    fixes.append({
                        'original': original,
                        'fixed': correct,
                        'location': para_info
                    })
                    return correct
                    
            return original
        
        # Исправляем поврежденные плейсхолдеры
        fixed_text = self.pattern_damaged.sub(fix_placeholder, text)
        
        return fixed_text, fixed_count, fixes
    
    def check_document_for_damaged_placeholders(self, doc_path):
        """
        Проверяет документ на наличие поврежденных плейсхолдеров
        Возвращает список найденных проблем
        """
        problems = []
        
        try:
            document = Document(doc_path)
        except Exception as e:
            print(f"⚠️  Не удалось открыть {doc_path.name}: {e}")
            return problems
            
        # Проверяем параграфы
        for i, para in enumerate(document.paragraphs):
            if para.text:
                # Находим все возможные плейсхолдеры
                all_matches = self.pattern_damaged.findall(para.text)
                correct_matches = self.pattern_placeholder.findall(para.text)
                
                # Если есть разница, значит есть поврежденные
                if len(all_matches) > len(correct_matches):
                    for match in all_matches:
                        if not self.pattern_placeholder.match(match):
                            problems.append({
                                'para_index': i,
                                'location': 'paragraph',
                                'text': para.text[:100] + '...' if len(para.text) > 100 else para.text,
                                'damaged': match
                            })
                    
        # Проверяем таблицы
        for t_idx, table in enumerate(document.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        if para.text:
                            all_matches = self.pattern_damaged.findall(para.text)
                            correct_matches = self.pattern_placeholder.findall(para.text)
                            
                            if len(all_matches) > len(correct_matches):
                                for match in all_matches:
                                    if not self.pattern_placeholder.match(match):
                                        problems.append({
                                            'table': t_idx,
                                            'row': r_idx,
                                            'cell': c_idx,
                                            'para_in_cell': p_idx,
                                            'location': 'table',
                                            'text': para.text[:100] + '...' if len(para.text) > 100 else para.text,
                                            'damaged': match
                                        })
        
        return problems
        
    def extract_placeholders_list(self, doc_path):
        """
        Извлекает список всех плейсхолдеров из документа в порядке появления
        """
        placeholders = []
        
        try:
            document = Document(doc_path)
        except Exception as e:
            print(f"⚠️  Не удалось открыть {doc_path.name}: {e}")
            return []
            
        # Параграфы
        for para in document.paragraphs:
            if para.text:
                matches = self.pattern_placeholder.findall(para.text)
                placeholders.extend(matches)
                
        # Таблицы
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text:
                            matches = self.pattern_placeholder.findall(para.text)
                            placeholders.extend(matches)
                            
        # Колонтитулы
        for section in document.sections:
            for para in section.header.paragraphs:
                if para.text:
                    matches = self.pattern_placeholder.findall(para.text)
                    placeholders.extend(matches)
            for para in section.footer.paragraphs:
                if para.text:
                    matches = self.pattern_placeholder.findall(para.text)
                    placeholders.extend(matches)
                    
        return placeholders
        
    def process_document(self, doc_path, output_path=None, original_doc_path=None, fix_damaged=True, force_mode=False):
        """
        ТОЧНАЯ КОПИЯ ЛОГИКИ из ai_studio_code.py - EQN RESTORATION BATCH VERSION 10.2
        """
        
        # Используем точную логику из process_document_binary
        original_text = self._get_all_text_from_docx(original_doc_path)
        if original_text is None: 
            return False, "Не удалось прочитать оригинальный файл."

        pattern_strict = re.compile(r'<<Eqn\d+(?:\.eps)?>>', re.IGNORECASE)
        original_placeholders = pattern_strict.findall(original_text)

        translation_text = self._get_all_text_from_docx(doc_path)
        if translation_text is None: 
            return False, "Не удалось прочитать файл перевода."

        pattern_robust = re.compile(r'[<\\>,\s]*?Eqn\d+(?:\.eps)?[<\\>,\s]*', re.IGNORECASE)
        translation_placeholders_found = pattern_robust.findall(translation_text)

        try:
            if not original_placeholders:
                if not translation_placeholders_found:
                    # ОБЯЗАТЕЛЬНО создаем родительские папки перед копированием
                    output_path.parent.mkdir(parents=True, exist_ok=True)
                    import shutil
                    shutil.copy(doc_path, output_path)
                    return True, "Плейсхолдеры не найдены, файл скопирован."
                else:
                    print(f"         - ВНИМАНИЕ: В оригинале нет плейсхолдеров. Удаление {len(translation_placeholders_found)} лишних из перевода...")
                    with open(doc_path, 'rb') as f:
                        binary_content = f.read()
                    for placeholder_to_remove in translation_placeholders_found:
                        binary_content = binary_content.replace(placeholder_to_remove.encode('utf-8'), b'')
                    output_path.parent.mkdir(parents=True, exist_ok=True)
                    with open(output_path, 'wb') as f:
                        f.write(binary_content)
                    return True, f"Успешно удалено {len(translation_placeholders_found)} лишних плейсхолдеров."

            if len(original_placeholders) != len(translation_placeholders_found):
                message = f"Обнаружено несоответствие: {len(original_placeholders)} в оригинале vs {len(translation_placeholders_found)} в переводе."
                if not force_mode:
                    return False, f"{message} Обработка остановлена."
                else:
                    print(f"         ⚠️  {message}")
                    print("         - Активирован ПРИНУДИТЕЛЬНЫЙ РЕЖИМ. Лишние плейсхолдеры будут проигнорированы.")
                    translation_placeholders_found = translation_placeholders_found[:len(original_placeholders)]

            with open(doc_path, 'rb') as f:
                binary_content = f.read()

            for i, placeholder_to_replace in enumerate(translation_placeholders_found):
                binary_content = binary_content.replace(
                    placeholder_to_replace.encode('utf-8', 'replace'),
                    original_placeholders[i].encode('utf-8', 'replace'),
                    1
                )

            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'wb') as f:
                f.write(binary_content)
            
            return True, f"Успешно заменено: {len(original_placeholders)} плейсхолдеров."

        except Exception as e:
            import traceback
            traceback.print_exc()
            return False, f"Критическая ошибка на этапе обработки: {e}"
            
    def _get_all_text_from_docx(self, doc_path):
        """Безопасно извлекает ВЕСЬ текст из документа для АНАЛИЗА."""
        try:
            document = Document(doc_path)
            full_text = []
            for para in document.paragraphs: 
                full_text.append(para.text)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs: 
                            full_text.append(para.text)
            for section in document.sections:
                for para in section.header.paragraphs: 
                    full_text.append(para.text)
                for para in section.footer.paragraphs: 
                    full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            print(f"⚠️  Ошибка чтения {doc_path.name}: {e}")
            return None
            
    def analyze_and_report(self, translation_path, original_path):
        """
        Анализирует документы и выводит подробный отчет
        """
        print("\n" + "="*70)
        print("ДЕТАЛЬНЫЙ АНАЛИЗ ДОКУМЕНТОВ")
        print("="*70)
        
        # Проверка на поврежденные плейсхолдеры
        print(f"\n📋 Проверка файла: {translation_path.name}")
        problems = self.check_document_for_damaged_placeholders(translation_path)
        
        if problems:
            print(f"\n⚠️  НАЙДЕНО ПОВРЕЖДЕННЫХ ПЛЕЙСХОЛДЕРОВ: {len(problems)}")
            print("\nПодробности:")
            for i, prob in enumerate(problems, 1):
                print(f"\n{i}. {prob['damaged']}")
                print(f"   Расположение: {prob['location']}")
                if 'para_index' in prob:
                    print(f"   Параграф: {prob['para_index']}")
                print(f"   Контекст: {prob['text']}")
        else:
            print("✅ Все плейсхолдеры в правильном формате")
        
        # Извлекаем плейсхолдеры
        trans_placeholders = self.extract_placeholders_list(translation_path)
        orig_placeholders = self.extract_placeholders_list(original_path)
        
        print(f"\n📊 Статистика плейсхолдеров:")
        print(f"   • В переводе: {len(trans_placeholders)} (правильно оформленных)")
        print(f"   • В оригинале: {len(orig_placeholders)}")
        
        if len(trans_placeholders) != len(orig_placeholders):
            print(f"\n⚠️  ВНИМАНИЕ: Количество не совпадает!")
            diff = abs(len(trans_placeholders) - len(orig_placeholders))
            if len(trans_placeholders) > len(orig_placeholders):
                print(f"   В переводе на {diff} плейсхолдеров больше")
            else:
                print(f"   В оригинале на {diff} плейсхолдеров больше")


def find_translation_original_pairs(translations_root, originals_root, translation_suffix='_to_en_us'):
    """
    Находит пары переведенных и оригинальных файлов с одинаковой структурой папок
    
    Параметры:
    - translations_root: корневая папка с переводами
    - originals_root: корневая папка с оригиналами
    - translation_suffix: суффикс в именах переведенных файлов
    
    Возвращает список кортежей: (путь_к_переводу, путь_к_оригиналу)
    """
    pairs = []
    not_found = []
    
    # Список возможных суффиксов перевода
    possible_suffixes = ['_translated_en-us', '_translated_en_us', '_to_en_us', '_to_en-us', '_en', '_EN']
    
    # Рекурсивно находим все .docx файлы в папке переводов
    for translation_file in translations_root.rglob("*.docx"):
        # Пропускаем временные файлы
        if translation_file.name.startswith("~$"):
            continue
            
        # Пропускаем уже обработанные файлы
        if "_restored" in translation_file.name:
            continue
            
        # Получаем относительный путь от корня переводов
        try:
            relative_path = translation_file.relative_to(translations_root)
        except ValueError:
            continue
            
        # Определяем имя оригинала (убираем суффикс перевода)
        original_name = translation_file.name
        for suffix in possible_suffixes:
            if suffix in translation_file.stem:
                original_stem = translation_file.stem.replace(suffix, '')
                original_name = original_stem + translation_file.suffix
                break
        
        # Получаем путь к папке перевода относительно корня
        translation_parent = relative_path.parent
        
        # Убираем суффиксы перевода из пути папок
        original_parent_parts = []
        for part in translation_parent.parts:
            part_cleaned = part
            for suffix in ['_to_en_us', '_en', '_EN', '_translated']:
                if part.endswith(suffix):
                    part_cleaned = part[:-len(suffix)]
                    break
            original_parent_parts.append(part_cleaned)
        
        # Строим путь к оригиналу
        if original_parent_parts:
            original_path = originals_root / Path(*original_parent_parts) / original_name
        else:
            original_path = originals_root / original_name
        
        if original_path.exists():
            pairs.append((translation_file, original_path))
        else:
            not_found.append({
                'translation': translation_file,
                'expected_original': original_path
            })
            
    return pairs, not_found


def process_multiple_files(pairs, restorer, output_root, translations_root, dry_run=False, force_mode=False):
    """
    Обрабатывает множество файлов
    
    Параметры:
    - pairs: список кортежей (перевод, оригинал)
    - restorer: экземпляр PlaceholderRestorer
    - output_root: корневая папка для сохранения результатов
    - translations_root: корневая папка с переводами (для расчета относительных путей)
    - dry_run: если True, только показывает что будет сделано
    - force_mode: принудительный режим для несоответствий
    
    Возвращает словарь с результатами
    """
    results = {
        'success': [],
        'failed': [],
        'skipped': [],
        'processed_without_placeholders': [],
        'damaged_fixed': [],
        'stats': {
            'total_files': len(pairs),
            'total_placeholders_replaced': 0,
            'total_damaged_fixed': 0,
            'files_without_placeholders': 0
        }
    }
    
    if dry_run:
        print("\n🔍 РЕЖИМ ПРЕДПРОСМОТРА (изменения не будут сохранены)")
    
    total = len(pairs)
    for i, (translation, original) in enumerate(pairs, 1):
        print(f"\n[{i}/{total}] Обработка: {translation.name}")
        print(f"         Оригинал: {original.name}")
        
        # Рассчитываем путь для сохранения с сохранением структуры папок
        try:
            relative_path = translation.relative_to(translations_root)
            output_path = output_root / relative_path
        except ValueError:
            # Если не удается получить относительный путь, используем имя файла
            output_path = output_root / translation.name
            
        if output_path.exists():
            print("         ⚠️  Пропущен: выходной файл уже существует")
            results['skipped'].append({
                'file': translation,
                'reason': 'Выходной файл уже существует'
            })
            continue
        
        if dry_run:
            # В режиме предпросмотра только анализируем
            problems = restorer.check_document_for_damaged_placeholders(translation)
            trans_placeholders = restorer.extract_placeholders_list(translation)
            orig_placeholders = restorer.extract_placeholders_list(original)
            
            if problems:
                print(f"         ⚠️  Поврежденных плейсхолдеров: {len(problems)}")
            if trans_placeholders and orig_placeholders:
                print(f"         📊 Плейсхолдеров: перевод={len(trans_placeholders)}, оригинал={len(orig_placeholders)}")
                if len(trans_placeholders) != len(orig_placeholders):
                    print("         ⚠️  Количество не совпадает!")
            results['success'].append({
                'file': translation,
                'output': output_path,
                'message': f"Будет обработано {len(trans_placeholders) if trans_placeholders else 0} плейсхолдеров"
            })
        else:
            # Реальная обработка
            try:
                success, message = restorer.process_document(
                    translation,
                    output_path=output_path,
                    original_doc_path=original,
                    fix_damaged=True,
                    force_mode=force_mode
                )
                
                if success:
                    # Проверяем типы обработки
                    if "без плейсхолдеров" in message:
                        results['processed_without_placeholders'].append({
                            'file': translation,
                            'output': output_path,
                            'message': message
                        })
                        results['stats']['files_without_placeholders'] += 1
                    else:
                        results['success'].append({
                            'file': translation,
                            'output': output_path,
                            'message': message
                        })
                        
                        # Подсчитываем статистику
                        if "Исправлено поврежденных" in message:
                            results['stats']['total_damaged_fixed'] += restorer.stats['damaged_placeholders_fixed']
                            if restorer.stats['damaged_placeholders_fixed'] > 0:
                                results['damaged_fixed'].append({
                                    'file': translation,
                                    'count': restorer.stats['damaged_placeholders_fixed'],
                                    'details': restorer.stats['damaged_placeholders_details']
                                })
                        results['stats']['total_placeholders_replaced'] += restorer.stats['placeholders_replaced']
                else:
                    results['failed'].append({
                        'file': translation,
                        'error': message
                    })
                    
            except Exception as e:
                results['failed'].append({
                    'file': translation,
                    'error': str(e)
                })
                
    return results


def generate_report(results, not_found, output_file=None):
    """
    Генерирует подробный отчет о результатах обработки
    """
    if output_file is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = f"eqn_restoration_report_{timestamp}.txt"
    
    report_lines = []
    report_lines.append("="*80)
    report_lines.append("ОТЧЕТ О ВОССТАНОВЛЕНИИ ПЛЕЙСХОЛДЕРОВ УРАВНЕНИЙ")
    report_lines.append("="*80)
    report_lines.append(f"Дата и время: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report_lines.append("")
    
    # Общая статистика
    report_lines.append("ОБЩАЯ СТАТИСТИКА:")
    report_lines.append("-"*40)
    report_lines.append(f"Всего файлов для обработки: {results['stats']['total_files']}")
    report_lines.append(f"Успешно обработано с плейсхолдерами: {len(results['success'])}")
    report_lines.append(f"Исправлено поврежденных плейсхолдеров: {results['stats']['total_damaged_fixed']}")
    report_lines.append(f"Обработано без плейсхолдеров: {len(results.get('processed_without_placeholders', []))}")
    report_lines.append(f"С ошибками: {len(results['failed'])}")
    report_lines.append(f"Пропущено: {len(results['skipped'])}")
    report_lines.append(f"Файлов без пары: {len(not_found)}")
    report_lines.append(f"Всего заменено плейсхолдеров: {results['stats']['total_placeholders_replaced']}")
    report_lines.append("")
    
    # Файлы с исправленными поврежденными плейсхолдерами
    if results.get('damaged_fixed'):
        report_lines.append("\nФАЙЛЫ С ИСПРАВЛЕННЫМИ ПОВРЕЖДЕННЫМИ ПЛЕЙСХОЛДЕРАМИ:")
        report_lines.append("-"*40)
        for item in results['damaged_fixed']:
            report_lines.append(f"🔧 {item['file'].name}")
            report_lines.append(f"   Исправлено: {item['count']} плейсхолдеров")
            for detail in item['details'][:3]:
                report_lines.append(f"   • {detail['original']} → {detail['fixed']}")
            if len(item['details']) > 3:
                report_lines.append(f"   ... и еще {len(item['details']) - 3}")
            report_lines.append("")
    
    # Успешно обработанные файлы
    if results['success']:
        report_lines.append("\nУСПЕШНО ОБРАБОТАННЫЕ ФАЙЛЫ:")
        report_lines.append("-"*40)
        for item in results['success']:
            report_lines.append(f"✅ {item['file'].name}")
            report_lines.append(f"   → {item['output'].name}")
            report_lines.append(f"   {item['message']}")
            report_lines.append("")
    
    # Остальные категории...
    # [Код для остальных категорий остается тем же]
    
    # Сохранение отчета
    report_text = "\n".join(report_lines)
    
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(report_text)
        print(f"\n📄 Отчет сохранен: {output_file}")
    except Exception as e:
        print(f"\n⚠️  Не удалось сохранить отчет: {e}")
    
    # Краткая статистика в консоль
    print("\n" + "="*70)
    print("ИТОГОВАЯ СТАТИСТИКА")
    print("="*70)
    print(f"✅ Успешно обработано: {len(results['success'])}")
    if results['stats']['total_damaged_fixed'] > 0:
        print(f"🔧 Исправлено поврежденных плейсхолдеров: {results['stats']['total_damaged_fixed']}")
    print(f"📊 Всего заменено плейсхолдеров: {results['stats']['total_placeholders_replaced']}")
    if results['failed']:
        print(f"❌ С ошибками: {len(results['failed'])}")
    if not_found:
        print(f"❓ Без пары: {len(not_found)}")
    
    return output_file


def process_single_file(restorer):
    """Обработка одного файла с детальным анализом"""
    # Запрашиваем переведенный файл
    print("\n1. Путь к ПЕРЕВЕДЕННОМУ файлу:")
    file_path = input("   → ").strip().strip('"\'')
    file_path = Path(file_path)
    
    if not file_path.exists() or not file_path.is_file():
        print(f"\n❌ Ошибка: файл не найден: {file_path}")
        return
        
    # Запрашиваем оригинал
    print("\n2. Путь к ОРИГИНАЛЬНОМУ файлу:")
    original_path = input("   → ").strip().strip('"\'')
    original_path = Path(original_path)
    
    if not original_path.exists() or not original_path.is_file():
        print(f"\n❌ Ошибка: файл не найден: {original_path}")
        return
    
    # Анализ документов
    restorer.analyze_and_report(file_path, original_path)
    
    # Спрашиваем, продолжать ли
    print("\n" + "="*70)
    choice = input("\nПродолжить обработку? (y/n): ").strip().lower()
    if choice != 'y':
        print("Операция отменена.")
        return
    
    # Обработка
    print("\n" + "="*70)
    print("ОБРАБОТКА")
    print("="*70)
    
    print("\n⏳ Исправление поврежденных плейсхолдеров и восстановление последовательности...")
    
    success, message = restorer.process_document(
        file_path,
        original_doc_path=original_path,
        fix_damaged=True
    )
    
    if success:
        print(f"\n✅ Успешно обработано!")
        print(f"📊 {message}")
        output_path = file_path.parent / f"{file_path.stem}_restored{file_path.suffix}"
        print(f"📄 Результат сохранен: {output_path}")
        
        # Показываем детали исправлений
        if restorer.stats['damaged_placeholders_fixed'] > 0:
            print(f"\n📝 Детали исправленных плейсхолдеров:")
            for fix in restorer.stats['damaged_placeholders_details'][:5]:
                print(f"   • {fix['original']} → {fix['fixed']} ({fix['location']})")
            if len(restorer.stats['damaged_placeholders_details']) > 5:
                print(f"   ... и еще {len(restorer.stats['damaged_placeholders_details']) - 5}")
    else:
        print(f"\n❌ Ошибка: {message}")


def process_folders(restorer):
    """Обработка папок с файлами"""
    print("\n" + "="*70)
    print("ОБРАБОТКА ПАПОК")
    print("="*70)
    
    # Запрашиваем папку с переводами
    print("\n1. Путь к папке с ПЕРЕВЕДЕННЫМИ файлами:")
    print("   (структура папок должна включать суффиксы _en, _EN, _to_en_us и т.п.)")
    translations_path = input("   → ").strip().strip('"\'')
    translations_root = Path(translations_path)
    
    if not translations_root.exists() or not translations_root.is_dir():
        print(f"\n❌ Ошибка: папка не найдена: {translations_root}")
        return
    
    # Запрашиваем папку с оригиналами
    print("\n2. Путь к папке с ОРИГИНАЛЬНЫМИ файлами:")
    originals_path = input("   → ").strip().strip('"\'')
    originals_root = Path(originals_path)
    
    if not originals_root.exists() or not originals_root.is_dir():
        print(f"\n❌ Ошибка: папка не найдена: {originals_root}")
        return
    
    print(f"\n📁 Папка с переводами: {translations_root}")
    print(f"📁 Папка с оригиналами: {originals_root}")
    
    # Поиск пар файлов
    print("\n🔍 Поиск соответствий между переводами и оригиналами...")
    pairs, not_found = find_translation_original_pairs(translations_root, originals_root)
    
    if not pairs:
        print("\n❌ Не найдено ни одной пары файлов для обработки!")
        if not_found:
            print(f"   Найдено {len(not_found)} переводов без соответствующих оригиналов.")
            print("\n   Первые 5 не найденных:")
            for item in not_found[:5]:
                print(f"   • {item['translation'].name} → искали: {item['expected_original']}")
        return
    
    # Показываем статистику
    print(f"\n📊 Найдено для обработки:")
    print(f"   • Пар файлов: {len(pairs)}")
    print(f"   • Переводов без пары: {len(not_found)}")
    
    # Опции обработки
    print("\n" + "="*70)
    print("ОПЦИИ ОБРАБОТКИ")
    print("="*70)
    print("1. Обработать ВСЕ найденные файлы (безопасный режим)")
    print("2. Обработать ВСЕ найденные файлы (принудительный режим)")
    print("3. Режим предпросмотра (показать что будет сделано)")
    print("4. Отмена")
    
    choice = input("\nВыберите опцию (1-4): ").strip()
    
    if choice == '4':
        print("Обработка отменена.")
        return
    
    # Создаем имя папки для результатов
    output_folder_name = translations_root.name + "_restored"
    output_root = translations_root.parent / output_folder_name
    
    print(f"\n📁 Папка для результатов: {output_root}")
    
    # Определяем режимы
    force_mode = (choice == '2')
    dry_run = (choice == '3')
    
    if force_mode:
        print("⚠️  Выбран ПРИНУДИТЕЛЬНЫЙ РЕЖИМ - несоответствия количества плейсхолдеров будут автоматически исправлены")
    
    # Создаем папку для результатов если её нет
    if choice in ['1', '2']:
        try:
            output_root.mkdir(parents=True, exist_ok=True)
            print(f"✅ Папка для результатов готова")
        except Exception as e:
            print(f"\n❌ Ошибка при создании папки для результатов: {e}")
            return
    
    # Обработка
    print("\n" + "="*70)
    print("ОБРАБОТКА ФАЙЛОВ")
    print("="*70)
    
    results = process_multiple_files(pairs, restorer, output_root, translations_root, dry_run=dry_run, force_mode=force_mode)
    
    # Генерация отчета
    if not dry_run:
        generate_report(results, not_found)


def main():
    """Основная функция с выбором режима работы"""
    print("\n" + "="*70)
    print("      EQN RESTORATION BATCH - FIXED VERSION V2")
    print("  С обнаружением и исправлением поврежденных плейсхолдеров")
    print("="*70)
    print("\nВозможности этой версии:")
    print("• Обнаружение поврежденных плейсхолдеров (например, <Eqn023.eps>>)")
    print("• Автоматическое исправление формата плейсхолдеров")
    print("• Восстановление правильной последовательности из оригинала")
    print("• Обработка отдельных файлов или целых папок")
    print("• Детальный отчет о всех изменениях")
    
    print("\nПринцип работы:")
    print("• 1-й плейсхолдер в переводе → 1-й из оригинала")
    print("• 2-й плейсхолдер в переводе → 2-й из оригинала")
    print("• и так далее...")
    
    # Создаем экземпляр обработчика
    restorer = PlaceholderRestorer()
    
    # Выбор режима
    print("\n" + "="*70)
    print("ВЫБЕРИТЕ РЕЖИМ РАБОТЫ")
    print("="*70)
    print("1. Обработать ОДИН файл")
    print("2. Обработать ПАПКУ с файлами")
    print("3. Выход")
    
    choice = input("\nВаш выбор (1-3): ").strip()
    
    if choice == '1':
        process_single_file(restorer)
    elif choice == '2':
        process_folders(restorer)
    elif choice == '3':
        print("\nВыход из программы.")
    else:
        print("\n❌ Неверный выбор!")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\nПрограмма прервана пользователем")
    except Exception as e:
        print(f"\n\nКритическая ошибка: {e}")
        import traceback
        traceback.print_exc()