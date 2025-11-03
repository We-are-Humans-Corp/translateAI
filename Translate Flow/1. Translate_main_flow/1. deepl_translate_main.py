import deepl
import os
import sys
import time
import traceback
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
import re  # Для регулярных выражений
from dotenv import load_dotenv

# === ДОБАВИТЬ ПОСЛЕ ВСЕХ ИМПОРТОВ (перед проверкой DEEPL_API_KEY) ===

import argparse


def parse_arguments():
    """Парсит аргументы командной строки для автоматического режима"""
    parser = argparse.ArgumentParser(description='DeepL Document Translator')
    parser.add_argument('--input', type=str, help='Путь к входному файлу')
    parser.add_argument('--output', type=str, help='Путь к выходному файлу')
    parser.add_argument('--source', type=str, default='RU', help='Исходный язык')
    parser.add_argument('--target', type=str, default='EN-US', help='Целевой язык')
    parser.add_argument('--no-interactive', action='store_true', help='Автоматический режим без запросов')

    return parser.parse_args()


# Импортируем словарь глоссария из локального файла
try:
    from deepl_glossary_python import glossary_entries
except ImportError:
    print("Предупреждение: Не удалось импортировать глоссарий из deepl_glossary_python.py")
    print("Перевод будет выполнен без использования глоссария.")
    glossary_entries = None

# --- Проверка и импорт python-docx ---
try:
    from docx import Document
    from docx.shared import Pt  # Для возможной работы со стилями, если понадобится
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
except ImportError:
    print("Ошибка: Необходима библиотека python-docx.")
    print("Пожалуйста, установите ее: pip install python-docx")
    sys.exit(1)
# --- Конец добавления ---

# --- Загрузка переменных окружения ---
# Определяем путь к .env файлу (теперь нужно подняться на 3 Translate_politics уровня вверх)
BASE_DIR = Path(__file__).resolve().parent.parent.parent
load_dotenv(BASE_DIR / '.env')

# --- Configuration ---
# Получаем ключ из переменных окружения
DEEPL_API_KEY = os.getenv('DEEPL_API_KEY', '')

SUPPORTED_EXTENSIONS = ['.docx', '.pdf']
TRANSLATION_SUFFIX = "_to_{target_lang_code}"
MAX_CONCURRENT_TRANSLATIONS = 1  # Всегда последовательный перевод
PAUSE_BETWEEN_REQUESTS = 2  # Пауза между запросами в секундах

LANGUAGE_OPTIONS = {
    "1": {"name": "Русский -> Английский (US)", "source": "RU", "target": "EN-US"},
    "2": {"name": "Английский -> Русский", "source": "EN", "target": "RU"},
    "3 Translate_politics": {"name": "Немецкий -> Русский", "source": "DE", "target": "RU"},
    # Добавьте другие языки при необходимости
}


# --- Helper Functions ---

def check_api_key_placeholder():
    """Checks if the API key is present in environment variables."""
    if not DEEPL_API_KEY:
        print("Ошибка: Ключ DeepL API не найден в переменных окружения.")
        print("Пожалуйста, создайте файл .env в корне проекта и добавьте:")
        print("DEEPL_API_KEY=your_api_key_here")
        print("\nИли установите переменную окружения DEEPL_API_KEY")
        sys.exit(1)


def get_translation_direction():
    """Запрашивает направление перевода у пользователя"""
    print("\nВыберите направление перевода:")
    for key, value in LANGUAGE_OPTIONS.items():
        print(f"  {key}: {value['name']}")

    while True:
        choice = input(f"Введите номер ({', '.join(LANGUAGE_OPTIONS.keys())}): ").strip()
        if choice in LANGUAGE_OPTIONS:
            selected_lang = LANGUAGE_OPTIONS[choice]
            print(f"Выбрано: {selected_lang['name']}")
            return selected_lang['source'], selected_lang['target']
        else:
            print("Ошибка: Неверный выбор. Пожалуйста, введите один из предложенных номеров.")


def initialize_translator(api_key):
    """Initializes the DeepL translator and checks usage."""
    print("\nИнициализация переводчика DeepL...")
    try:
        translator = deepl.Translator(api_key)
        print("Проверка лимитов DeepL API...")
        usage = translator.get_usage()

        # Character limit check
        char_count = usage.character.count
        char_limit = usage.character.limit
        limit_reached_char = False

        # Калькулятор затрат (€20.00 за 1 миллион символов)
        PRICE_PER_MILLION = 20.00  # EUR

        if char_count is not None:
            # Рассчитываем затраты
            millions_used = char_count / 1_000_000
            total_cost = millions_used * PRICE_PER_MILLION

            # Форматирование для отображения
            count_str = f"{char_count:,}"
            cost_str = f"€{total_cost:.2f}"

            print(f"\n{'=' * 60}")
            print(f"💰 КАЛЬКУЛЯТОР ЗАТРАТ DeepL API")
            print(f"{'=' * 60}")
            print(f"📊 Использовано символов: {count_str}")
            print(f"💶 Тариф: €{PRICE_PER_MILLION:.2f} за 1 миллион символов")
            print(f"💸 ОБЩИЕ ЗАТРАТЫ: {cost_str}")
            print(f"{'=' * 60}")

            if char_limit is not None:
                limit_str = f"{char_limit:,}"
                remaining = max(0, char_limit - char_count)
                remaining_str = f"{remaining:,}"
                percentage_used = (char_count / char_limit) * 100

                print(f"\n📈 Статистика использования:")
                print(f"  • Лимит: {limit_str} символов")
                print(f"  • Использовано: {count_str} ({percentage_used:.2f}%)")
                print(f"  • Осталось: {remaining_str}")

                if char_count >= char_limit:
                    limit_reached_char = True
                    print("  ⚠️ ВНИМАНИЕ: Лимит символов исчерпан!")
            else:
                print(f"\n📈 Статистика использования:")
                print(f"  • Использовано: {count_str} символов")
                print(f"  • Лимит: не установлен")
        else:
            print("  Лимит символов: Использование не отслеживается или не применимо.")

        # Document limit check
        doc_count = None
        doc_limit = None
        limit_reached_doc = False
        if hasattr(usage, 'document') and usage.document and usage.document.valid:
            doc_count = usage.document.count
            doc_limit = usage.document.limit
            if doc_limit is not None:
                count_str = f"{doc_count:,}" if doc_count is not None else "N/A"
                limit_str = f"{doc_limit:,}"
                remaining_str = ""
                if doc_count is not None:
                    remaining_str = f"(Осталось: {max(0, doc_limit - doc_count):,})"
                    if doc_count is not None and doc_limit is not None and doc_count >= doc_limit: limit_reached_doc = True
                print(f"  Лимит документов: Использовано {count_str} из {limit_str} {remaining_str}")
                if limit_reached_doc: print("Внимание: Лимит документов для перевода исчерпан или достигнут!")
            elif doc_count is not None:
                print(f"  Лимит документов: Использовано {doc_count:,} (Лимит не установлен или не известен)")
            else:
                print("  Лимит документов: Использование не отслеживается или не применимо для этого плана.")
        else:
            print("  Лимит документов: Не применимо или не предоставлено API (возможно, DeepL API Free?)")

        # Final check if any limit is exceeded
        if limit_reached_char or limit_reached_doc:
            error_msg = "Достигнут лимит использования DeepL API."
            if limit_reached_char: error_msg = "Достигнут лимит символов DeepL API."
            if limit_reached_doc: error_msg = "Достигнут лимит документов DeepL API."
            print(f"\nОшибка: {error_msg}")
            print("Перевод невозможен.")
            sys.exit(1)
        else:
            print("Инициализация DeepL успешна. Лимиты в норме.")
        return translator

    except deepl.AuthorizationException:
        print("\nОшибка аутентификации DeepL: Неверный API ключ.")
        sys.exit(1)
    except Exception as e:
        print(f"\nНепредвиденная ошибка при инициализации DeepL: {e}")
        print(traceback.format_exc())
        sys.exit(1)


def get_or_create_glossary(translator, source_lang, target_lang):
    """Получает существующий глоссарий или создает новый."""
    # Если глоссарий не был импортирован, возвращаем None
    if glossary_entries is None:
        return None

    # Проверяем, что направление перевода поддерживается глоссарием
    if source_lang != "RU" or target_lang not in ["EN-US", "EN-GB", "EN"]:
        print("Глоссарий доступен только для перевода с русского на английский.")
        return None

    glossary_name = f"Universal Scientific Terms RU-EN v1"

    try:
        # Пытаемся найти существующий глоссарий
        print("\nПроверка существующих глоссариев...")
        glossaries = translator.list_glossaries()

        for glossary in glossaries:
            if glossary.name == glossary_name and glossary.source_lang == "RU" and glossary.target_lang == "EN":
                print(f"✅ Найден существующий глоссарий: {glossary.name} (ID: {glossary.glossary_id})")
                print(f"   Количество терминов: {glossary.entry_count}")
                return glossary

        # Если глоссарий не найден, создаем новый
        print("Создание нового глоссария научных терминов...")
        glossary = translator.create_glossary(
            name=glossary_name,
            source_lang="RU",
            target_lang="EN",
            entries=glossary_entries
        )
        print(f"✅ Глоссарий успешно создан: {glossary.name} (ID: {glossary.glossary_id})")
        print(f"   Количество терминов: {len(glossary_entries)}")
        return glossary

    except deepl.DeepLException as e:
        print(f"⚠️ Не удалось создать/получить глоссарий: {e}")
        print("   Перевод будет выполнен без использования глоссария.")
        return None
    except Exception as e:
        print(f"⚠️ Неожиданная ошибка при работе с глоссарием: {e}")
        print("   Перевод будет выполнен без использования глоссария.")
        return None


def find_files_to_translate(source_root_path):
    """Recursively finds files with supported extensions, excluding temporary files."""
    print("\nПоиск файлов для перевода...")
    files_found = []
    for extension in SUPPORTED_EXTENSIONS:
        extension_pattern = "".join([f"[{c.lower()}{c.upper()}]" for c in extension[1:]])
        pattern = f"*.{extension_pattern}"
        # print(f"Поиск файлов с шаблоном: {pattern} в {source_root_path}") # Debug log
        files_found.extend(list(source_root_path.rglob(pattern)))

    processed_paths = set()
    files_to_process = []
    for f in files_found:
        try:
            resolved_f = f.resolve()
            if resolved_f.is_file() and not resolved_f.name.startswith('~$') and resolved_f not in processed_paths:
                files_to_process.append(resolved_f)
                processed_paths.add(resolved_f)
        except Exception as e:
            print(f"Предупреждение: Не удалось обработать путь файла '{f}': {e}. Пропуск.")

    files_to_process.sort()
    print(f"Найдено {len(files_to_process)} файлов с расширениями {SUPPORTED_EXTENSIONS} (исключая временные файлы).")
    return files_to_process


def estimate_translation_cost(file_paths):
    """Оценивает примерную стоимость перевода файлов"""
    total_chars = 0
    file_estimates = []

    print("\n📊 Оценка стоимости перевода...")
    print("-" * 60)

    for file_path in file_paths[:10]:  # Показываем первые 10 файлов
        try:
            # Грубая оценка: ~2000 символов на страницу для DOCX, ~3000 для PDF
            file_size = file_path.stat().st_size

            if file_path.suffix.lower() == '.docx':
                # Примерная оценка: размер файла * 0.5 (очень грубо)
                estimated_chars = int(file_size * 0.5)
            elif file_path.suffix.lower() == '.pdf':
                # Примерная оценка: размер файла * 0.3 Translate_politics (очень грубо)
                estimated_chars = int(file_size * 0.3)
            else:
                estimated_chars = int(file_size * 0.4)

            total_chars += estimated_chars
            file_estimates.append((file_path.name, estimated_chars))

        except Exception:
            continue

    # Экстраполируем на все файлы, если их больше 10
    if len(file_paths) > 10:
        avg_chars = total_chars / min(10, len(file_paths))
        total_chars = int(avg_chars * len(file_paths))
        print(f"Анализ первых 10 файлов из {len(file_paths)}...")

    # Расчет стоимости
    PRICE_PER_MILLION = 20.00  # EUR
    millions = total_chars / 1_000_000
    estimated_cost = millions * PRICE_PER_MILLION

    print(f"\n💰 ОЦЕНКА СТОИМОСТИ:")
    print(f"  • Файлов для перевода: {len(file_paths)}")
    print(f"  • Примерное количество символов: {total_chars:,}")
    print(f"  • Тариф: €{PRICE_PER_MILLION:.2f} за 1 млн символов")
    print(f"  • ПРИМЕРНАЯ СТОИМОСТЬ: €{estimated_cost:.2f}")
    print("-" * 60)
    print("⚠️  Это грубая оценка. Реальная стоимость зависит от")
    print("   фактического содержания документов.")

    return total_chars, estimated_cost


def clean_translated_docx(docx_path):
    """
    Открывает переведенный DOCX и исправляет искаженные плейсхолдеры
    к виду <<EqnXXX.eps>>. Обрабатывает случаи с лишней точкой,
    отсутствующими >>, и отсутствующим .eps.
    Перезаписывает файл. Возвращает True при успехе/отсутствии изменений, False при ошибке.
    """
    try:
        if not docx_path.is_file():
            print(f"  -> ОШИБКА Пост-обработки: Файл не найден {docx_path}")
            return False

        document = Document(docx_path)
        changes_made = 0  # Счетчик изменений для логирования

        # УЛУЧШЕННОЕ РЕГУЛЯРНОЕ ВЫРАЖЕНИЕ v3:
        pattern = re.compile(
            r'(<<Eqn(\d+))'  # Группа 1: <<Eqn<цифры>, Группа 2: только цифры
            r'(?!'  # Начало негативного просмотра вперед (убеждаемся, что ДАЛЬШЕ НЕ...)
            r'\.eps>>'  # ...ровно ".eps>>"
            r'([,\s]|$)'  # ...за которым идет запятая, пробел или конец строки/параграфа
            r')'  # Конец негативного просмотра
            r'([\.\w>]+)?'  # Группа 3 Translate_politics (опциональная): Захватываем сам "мусор" - точки, буквы(eps), >.
        )

        # Строка для замены: Восстанавливаем правильный формат, используя Группу 2 (цифры)
        replacement_pattern = r'<<Eqn\2.eps>>'

        # Вспомогательная функция для обработки параграфа
        def process_paragraph(para: Paragraph):
            nonlocal changes_made
            # Быстрая проверка для оптимизации - ищем хотя бы начало плейсхолдера
            if '<<Eqn' not in para.text:
                return False

            original_text = para.text
            # Применяем замену ко всему тексту параграфа
            new_text = pattern.sub(replacement_pattern, original_text)

            if new_text != original_text:
                # Считаем количество реальных замен (хотя бы примерно)
                # Это не идеально точно, но дает представление
                num_replacements = len(pattern.findall(original_text))
                changes_made += num_replacements

                # Простой способ обновления: очистить параграф и вставить новый текст
                para.clear()  # Очищает все runs и текст
                para.add_run(new_text)  # Добавляет новый текст одним run (может сбросить форматирование)

                # Debug Log (можно раскомментировать для отладки)
                # print(f"    Debug: Replaced in para. Orig: '{original_text}'. New: '{new_text}'")
                return True
            return False

        # --- Основная логика обхода документа ---
        # Итерация по параграфам в основном тексте
        for para in document.paragraphs: process_paragraph(para)

        # Итерация по таблицам
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs: process_paragraph(para)

        # Итерация по колонтитулам (если нужно)
        for section in document.sections:
            # Headers
            for para in section.header.paragraphs: process_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs: process_paragraph(para)
            # Footers
            for para in section.footer.paragraphs: process_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs: process_paragraph(para)
        # --- Конец обхода ---

        if changes_made > 0:
            print(f"  -> Пост-обработка V3: Исправлено ~{changes_made} искаженных плейсхолдеров в {docx_path.name}")
            document.save(docx_path)  # Перезаписываем файл
        else:
            print(
                f"  -> Пост-обработка V3: Искаженные плейсхолдеры (требующие исправления) не найдены в {docx_path.name}")

        return True

    except FileNotFoundError:
        print(f"  -> ОШИБКА Пост-обработки: Файл не найден {docx_path}")
        return False
    except Exception as e:
        # Используем f-string для форматирования
        print(f"  -> КРИТИЧЕСКАЯ ОШИБКА Пост-обработки: Не удалось обработать {docx_path.name}: {e}")
        print(traceback.format_exc())  # Для отладки
        return False


def translate_single_document(input_path, output_path, translator, target_lang, source_lang, file_index, total_files,
                              source_root_path, glossary=None):
    """Translates a single document and handles errors. Includes enhanced post-processing for DOCX."""
    # Конвертируем Path объекты в строки для DeepL API и логгирования
    input_path_str = str(input_path)
    output_path_str = str(output_path)

    try:
        relative_path_for_display = input_path.relative_to(source_root_path)
    except ValueError:
        relative_path_for_display = input_path
    except Exception as e:
        relative_path_for_display = f"Ошибка при определении пути: {input_path_str}"  # Используем str
        print(f"Предупреждение: Не удалось определить относительный путь для отображения: {e}")

    # Используем имена файлов из Path объектов для большей точности
    print(f"\n[{file_index}/{total_files}] Обработка файла: {relative_path_for_display}")
    print(f"  Перевод '{input_path.name}' -> '{output_path.name}'...")
    print(f"     (Языки: {source_lang} -> {target_lang})")

    post_processing_error_occurred = False  # Флаг для отслеживания ошибок пост-обработки

    try:
        # Используем Path объект для создания директории
        output_path.parent.mkdir(parents=True, exist_ok=True)

        start_time = time.time()

        # Добавляем логирование использования глоссария
        if glossary:
            print(f"     (Используется глоссарий: {glossary.name})")

        # Передаем глоссарий в функцию перевода, если он доступен
        translation_kwargs = {
            "input_path": input_path_str,  # DeepL требует строку
            "output_path": output_path_str,  # DeepL требует строку
            "target_lang": target_lang,
            "source_lang": source_lang
        }

        if glossary:
            translation_kwargs["glossary"] = glossary

        translator.translate_document_from_filepath(**translation_kwargs)

        end_time = time.time()
        print(f"  -> УСПЕШНО переведен за {end_time - start_time:.2f} сек.")

        # Добавляем паузу между запросами
        print(f"  -> Пауза {PAUSE_BETWEEN_REQUESTS} секунды перед следующим запросом...")
        time.sleep(PAUSE_BETWEEN_REQUESTS)

        # --- ИНТЕГРАЦИЯ ПОСТ-ОБРАБОТКИ ---
        if output_path.suffix.lower() == '.docx':
            print(f"  -> Запуск пост-обработки V3 для {output_path.name}...")
            # Передаем Path объект в функцию очистки
            cleaning_successful = clean_translated_docx(output_path)
            if not cleaning_successful:
                # Логируем предупреждение и устанавливаем флаг
                print(f"  -> ПРЕДУПРЕЖДЕНИЕ: Пост-обработка файла {output_path.name} завершилась с ошибкой.")
                post_processing_error_occurred = True  # Отмечаем ошибку пост-обработки
        # --- КОНЕЦ ИНТЕГРАЦИИ ---

        # Возвращаем статус и пути. Добавляем информацию об ошибке пост-обработки.
        return {
            "status": "success" if not post_processing_error_occurred else "success_with_postprocessing_error",
            "input": input_path_str,
            "output": output_path_str,
            "post_processing_error": post_processing_error_occurred  # Явно указываем на проблему
        }

    # Обработка ошибок DeepL API и файловой системы
    except deepl.DocumentTranslationException as e:
        error_msg = f"Ошибка перевода документа DeepL: {e}"
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg}
    except deepl.QuotaExceededException:
        error_msg = "Превышена квота DeepL API."
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg, "quota_exceeded": True}
    except deepl.TooManyRequestsException:
        error_msg = "Слишком много запросов к DeepL API."
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg, "rate_limited": True}
    except deepl.ConnectionException as e:
        error_msg = f"Ошибка сети при обращении к DeepL: {e}"
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg}
    except deepl.DeepLException as e:  # Общая ошибка DeepL
        error_msg = f"Общая ошибка DeepL API: {e}"
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg}
    except FileNotFoundError:
        # Скорее всего, не найден ИСХОДНЫЙ файл
        error_msg = f"Исходный файл не найден: {input_path_str}"
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg}
    except PermissionError:
        error_msg = f"Нет прав на чтение/запись файла ({input_path_str} -> {output_path_str})."
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg}
    except Exception as e:
        # Перехват других непредвиденных ошибок (включая возможные ошибки при вызове clean_translated_docx)
        error_msg = f"Непредвиденная ошибка при обработке файла '{input_path_str}': {e}"
        print(f"  -> КРИТИЧЕСКАЯ ОШИБКА: {error_msg}")
        print(traceback.format_exc())
        return {"status": "error", "input": input_path_str, "output": output_path_str,
                "error": f"{error_msg}\n{traceback.format_exc()}"}


def process_translations(translator, file_paths, source_root_path, target_root_path, source_lang, target_lang,
                         suffix_pattern, glossary=None):
    """Manages the translation process using sequential processing."""
    success_count = 0
    success_with_warnings = 0  # Отдельный счетчик для успехов с ошибками пост-обработки
    skipped_suffix_count = 0
    skipped_exists_count = 0
    error_count = 0
    errors_list = []  # Список деталей ошибок
    total_files = len(file_paths)
    stop_processing = False  # Флаг для остановки при критических ошибках API

    actual_suffix = suffix_pattern.format(target_lang_code=target_lang.lower().replace("-", "_"))

    # Всегда используем 1 поток для последовательной обработки
    max_workers = 1
    processing_mode = "ПОСЛЕДОВАТЕЛЬНОЙ"

    print(f"\nЗапуск {processing_mode} обработки {total_files} файлов...")
    print("-" * 30)

    executor = ThreadPoolExecutor(max_workers=max_workers)
    future_to_path_obj = {}  # Храним Path объекты
    files_submitted = 0

    for i, input_path_obj in enumerate(file_paths):  # Используем Path объект
        if stop_processing:
            print(
                f"\n[{i + 1}/{total_files}] Пропуск отправки файла {input_path_obj.name} из-за критической ошибки API.")
            continue

        file_index = i + 1
        try:
            relative_path_for_target = input_path_obj.relative_to(source_root_path)
            output_path_obj = target_root_path / relative_path_for_target.with_name(
                f"{input_path_obj.stem}{actual_suffix}{input_path_obj.suffix}")
        except Exception as e:
            print(
                f"\n[{file_index}/{total_files}] Ошибка при расчете выходного пути для {input_path_obj.name}: {e}. Пропуск файла.")
            error_count += 1
            # Добавляем детали ошибки
            errors_list.append(
                {"status": "error", "file": str(input_path_obj), "error": f"Ошибка при расчете выходного пути: {e}"})
            continue

        # --- Skip Checks ---
        if input_path_obj.stem.endswith(actual_suffix):
            try:
                display_path = input_path_obj.relative_to(source_root_path)
            except ValueError:
                display_path = input_path_obj.name
            print(f"\n[{file_index}/{total_files}] Пропуск (файл уже имеет суффикс '{actual_suffix}'): {display_path}")
            skipped_suffix_count += 1
            continue

        if output_path_obj.exists():
            try:
                display_path = output_path_obj.relative_to(target_root_path)
            except ValueError:
                display_path = output_path_obj.name
            print(f"\n[{file_index}/{total_files}] Пропуск (переведенный файл уже существует): {display_path}")
            skipped_exists_count += 1
            continue
        # --- End Skip Checks ---

        # Submit the translation task
        future = executor.submit(
            translate_single_document,
            input_path_obj,  # Передаем Path
            output_path_obj,  # Передаем Path
            translator,
            target_lang,
            source_lang,
            file_index,
            total_files,
            source_root_path,  # Path
            glossary  # Передаем глоссарий
        )
        future_to_path_obj[future] = input_path_obj  # Храним Path объект
        files_submitted += 1

    print(f"\nОтправлено {files_submitted} заданий на перевод. Ожидание завершения...")

    # --- Сбор результатов ---
    try:
        for future in as_completed(future_to_path_obj):
            input_path_obj = future_to_path_obj[future]
            input_filename_str = str(input_path_obj)  # Строка для логов
            try:
                result = future.result()  # Получаем результат из потока

                if result['status'] == "success":
                    success_count += 1
                elif result['status'] == "success_with_postprocessing_error":
                    # Успешный перевод, но проблема с очисткой
                    success_with_warnings += 1
                    # Добавляем информацию об ошибке пост-обработки в общий список ошибок
                    errors_list.append({
                        "status": "warning",  # Используем статус warning
                        "file": result.get('input', input_filename_str),
                        "output": result.get('output'),
                        "error": f"Ошибка пост-обработки файла {Path(result.get('output', '')).name}"
                    })
                elif result['status'] == "error":
                    # Ошибка перевода или другая критическая ошибка
                    error_count += 1
                    # Убедимся, что ключ 'file' есть в словаре ошибки для отчета
                    if 'file' not in result: result['file'] = result.get('input', input_filename_str)
                    errors_list.append(result)
                    # Проверяем на критические ошибки API для остановки
                    if result.get("quota_exceeded") or result.get("rate_limited"):
                        if not stop_processing:
                            print("\n*** Обнаружено превышение квоты или лимита запросов DeepL. ***")
                            print("*** Новые задания не будут отправляться. Дождитесь завершения текущих. ***")
                        stop_processing = True
                else:
                    # Непредвиденный статус
                    print(
                        f"Предупреждение: Неизвестный статус результата '{result.get('status')}' для файла {input_filename_str}")
                    error_count += 1  # Считаем как ошибку
                    errors_list.append({"status": "error", "file": input_filename_str,
                                        "error": f"Неизвестный статус результата: {result.get('status')}"})


            except Exception as exc:
                # Ошибка получения результата из самого future (редко)
                error_count += 1
                error_msg = f"Критическая ошибка при получении результата для файла '{input_filename_str}': {exc}"
                print(f"  -> КРИТИЧЕСКАЯ ОШИБКА ПОТОКА: {error_msg}")
                print(traceback.format_exc())
                errors_list.append(
                    {"status": "error", "file": input_filename_str, "error": f"{error_msg}\n{traceback.format_exc()}"})
    finally:
        # Гарантированно закрываем пул потоков
        executor.shutdown(wait=True)

    print("-" * 30)
    print("Обработка завершена.")
    # Возвращаем все счетчики
    return success_count, success_with_warnings, skipped_suffix_count, skipped_exists_count, error_count, errors_list


# --- ФУНКЦИИ ДЛЯ РЕЖИМА ОДНОГО ФАЙЛА ---

def handle_single_file(translator):
    """Обработка перевода одного файла"""
    print("\n--- Режим перевода одного файла ---")

    # 1. Запрашиваем путь к файлу
    while True:
        source_str = input("\nВведите ПОЛНЫЙ путь к файлу для перевода: ").strip().strip('"\'')
        source_path = Path(source_str)
        if source_path.is_file() and source_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            break
        else:
            print(f"Ошибка: Файл не найден или не поддерживается.")
            print(f"Поддерживаемые форматы: {', '.join(SUPPORTED_EXTENSIONS)}")

    # 2. Запрашиваем направление перевода
    source_lang, target_lang = get_translation_direction()

    # 3 Translate_politics. Получаем или создаем глоссарий
    glossary = get_or_create_glossary(translator, source_lang, target_lang)

    # 4. Определяем имя выходного файла
    target_lang_code_for_suffix = target_lang.lower().replace("-", "_")
    actual_suffix = TRANSLATION_SUFFIX.format(target_lang_code=target_lang_code_for_suffix)
    output_path = source_path.with_name(f"{source_path.stem}{actual_suffix}{source_path.suffix}")

    print("-" * 50)
    print(f"Исходный файл: {source_path}")
    print(f"Целевой файл: {output_path}")
    print(f"Направление: {source_lang} -> {target_lang}")
    if glossary:
        print(f"Глоссарий: {glossary.name}")
    print("-" * 50)

    # 5. Проверяем существование файла
    if output_path.exists():
        overwrite = input(f"\nФайл '{output_path.name}' уже существует. Перезаписать? (y/n): ").strip().lower()
        if overwrite != 'y':
            print("Перевод отменен.")
            return

    # 6. Выполняем перевод
    result = translate_single_document(
        input_path=source_path,
        output_path=output_path,
        translator=translator,
        target_lang=target_lang,
        source_lang=source_lang,
        file_index=1,
        total_files=1,
        source_root_path=source_path.parent,
        glossary=glossary
    )

    # 7. Выводим результат
    print("\n" + "=" * 50)
    if result['status'] == "success":
        print("✅ Файл успешно переведен!")
        print(f"Сохранен как: {output_path}")
    elif result['status'] == "success_with_postprocessing_error":
        print("⚠️ Файл переведен, но возникла ошибка при пост-обработке.")
        print(f"Сохранен как: {output_path}")
    else:
        print("❌ Ошибка при переводе файла.")
        if 'error' in result:
            print(f"Детали: {result['error']}")
    print("=" * 50)


def find_largest_files():
    """Найти 10 самых больших файлов по количеству символов"""
    print("\n--- Поиск 10 самых больших файлов ---")
    print("\nВведите ПОЛНЫЙ путь к папке для поиска")
    print("Пример для Mac: /Users/имя/Desktop/МоиДокументы")
    print("Пример для Windows: C:\\Users\\имя\\Documents\\МоиДокументы")
    
    while True:
        search_path_str = input("-> ").strip().strip('"\'')
        search_path = Path(search_path_str)
        if search_path.is_dir():
            break
        else:
            print(f"Ошибка: Папка не найдена: '{search_path_str}'")
            print("Пожалуйста, введите корректный путь.")
    
    print(f"\nПоиск файлов в: {search_path}")
    print("Анализ размеров файлов...")
    
    # Собираем информацию о файлах
    file_data = []
    total_files = 0
    errors = 0
    
    # Поддерживаемые расширения для анализа
    extensions_to_check = ['.docx', '.pdf', '.txt', '.doc', '.rtf', '.odt']
    
    for ext in extensions_to_check:
        pattern = f"*{ext}"
        for file_path in search_path.rglob(pattern):
            total_files += 1
            try:
                if file_path.is_file() and not file_path.name.startswith('~$'):
                    # Получаем размер файла
                    file_size = file_path.stat().st_size
                    
                    # Оценка количества символов по типу файла
                    if ext in ['.txt']:
                        # Для текстовых файлов - примерно равно размеру
                        estimated_chars = file_size
                    elif ext in ['.docx', '.doc', '.odt']:
                        # Для документов Word - примерная оценка
                        estimated_chars = int(file_size * 0.7)
                    elif ext in ['.pdf']:
                        # Для PDF - очень грубая оценка
                        estimated_chars = int(file_size * 0.4)
                    elif ext in ['.rtf']:
                        # Для RTF - содержит много разметки
                        estimated_chars = int(file_size * 0.3)
                    else:
                        estimated_chars = int(file_size * 0.5)
                    
                    file_data.append({
                        'path': file_path,
                        'size': file_size,
                        'chars': estimated_chars,
                        'type': ext
                    })
            except Exception as e:
                errors += 1
                print(f"Ошибка при анализе файла {file_path.name}: {e}")
    
    if not file_data:
        print(f"\nНе найдено файлов для анализа в указанной папке.")
        print(f"Поддерживаемые форматы: {', '.join(extensions_to_check)}")
        return
    
    # Сортируем по количеству символов (убывание)
    file_data.sort(key=lambda x: x['chars'], reverse=True)
    
    # Берем топ-10
    top_files = file_data[:10]
    
    print(f"\n{'=' * 80}")
    print(f"ТОП-10 САМЫХ БОЛЬШИХ ФАЙЛОВ ПО КОЛИЧЕСТВУ СИМВОЛОВ")
    print(f"{'=' * 80}")
    print(f"Всего проанализировано файлов: {len(file_data)}")
    if errors > 0:
        print(f"Ошибок при анализе: {errors}")
    print(f"{'=' * 80}")
    
    for idx, file_info in enumerate(top_files, 1):
        try:
            # Относительный путь для удобства
            rel_path = file_info['path'].relative_to(search_path)
        except ValueError:
            rel_path = file_info['path'].name
        
        # Форматирование размеров
        size_mb = file_info['size'] / (1024 * 1024)
        chars_millions = file_info['chars'] / 1_000_000
        
        print(f"\n{idx}. {rel_path}")
        print(f"   Тип: {file_info['type']}")
        print(f"   Размер файла: {size_mb:.2f} MB ({file_info['size']:,} байт)")
        print(f"   Примерное кол-во символов: {file_info['chars']:,} (~{chars_millions:.2f} млн)")
        
        # Оценка стоимости перевода для этого файла
        cost_per_file = chars_millions * 20.00  # €20 за миллион
        print(f"   Примерная стоимость перевода: €{cost_per_file:.2f}")
    
    # Общая статистика по топ-10
    total_chars = sum(f['chars'] for f in top_files)
    total_size = sum(f['size'] for f in top_files)
    total_cost = (total_chars / 1_000_000) * 20.00
    
    print(f"\n{'=' * 80}")
    print(f"ИТОГО ДЛЯ ТОП-10 ФАЙЛОВ:")
    print(f"  Общий размер: {total_size / (1024 * 1024):.2f} MB")
    print(f"  Общее кол-во символов: {total_chars:,} (~{total_chars / 1_000_000:.2f} млн)")
    print(f"  Примерная стоимость перевода всех 10 файлов: €{total_cost:.2f}")
    print(f"{'=' * 80}")


def handle_folder_translation(translator):
    """Обработка папки с файлами"""
    print("\n--- Режим перевода папки ---")
    print("\nВведите ПОЛНЫЙ путь к папке с файлами для перевода")
    print("Пример для Mac: /Users/имя/Desktop/МоиДокументы")
    print("Пример для Windows: C:\\Users\\имя\\Documents\\МоиДокументы")

    while True:
        source_root_str = input("-> ").strip().strip('"\'')
        source_root_path = Path(source_root_str)
        if source_root_path.is_dir():
            break
        else:
            print(f"Ошибка: Папка не найдена или не является папкой: '{source_root_str}'")
            print("Пожалуйста, введите корректный путь.")

    # Запрашиваем направление перевода
    source_lang, target_lang = get_translation_direction()

    target_lang_code_for_suffix = target_lang.lower().replace("-", "_")
    actual_suffix = TRANSLATION_SUFFIX.format(target_lang_code=target_lang_code_for_suffix)
    target_root_path = source_root_path.with_name(f"{source_root_path.name}{actual_suffix}")

    print("-" * 30)
    print(f"Исходная папка: {source_root_path}")
    print(f"Целевая папка для переводов: {target_root_path}")
    print(f"Направление перевода: {source_lang} -> {target_lang}")
    print(f"Суффикс переведенных файлов: {actual_suffix}")
    print(f"Режим обработки: ПОСЛЕДОВАТЕЛЬНЫЙ (1 файл за раз)")
    print("-" * 30)

    print(f"\nПроверка/создание папки для переводов: {target_root_path}")
    try:
        target_root_path.mkdir(parents=True, exist_ok=True)
        print("Папка для переводов готова.")
    except PermissionError:
        print(f"Ошибка: Нет прав на создание папки '{target_root_path}'.")
        return
    except Exception as e:
        print(f"Ошибка при создании папки для переводов: {e}")
        print(traceback.format_exc())
        return

    files_to_translate = find_files_to_translate(source_root_path)

    if not files_to_translate:
        print("\nНе найдено подходящих файлов для перевода в указанной папке и ее подпапках.")
        return

    # Оценка стоимости перевода
    estimated_chars, estimated_cost = estimate_translation_cost(files_to_translate)

    # Запрос подтверждения
    print(f"\n⚠️  Продолжить перевод {len(files_to_translate)} файлов?")
    print(f"   Примерная стоимость: €{estimated_cost:.2f}")
    confirm = input("\nВведите 'y' для продолжения или любую другую клавишу для отмены: ").strip().lower()
    if confirm != 'y':
        print("Перевод отменен.")
        return

    # Получаем или создаем глоссарий (если применимо)
    glossary = get_or_create_glossary(translator, source_lang, target_lang)

    # Запоминаем начальное использование
    start_usage = translator.get_usage()
    start_chars = start_usage.character.count if start_usage.character.count else 0

    # Запуск процесса перевода и очистки
    success, warnings, skipped_suffix, skipped_exists, errors, error_details = process_translations(
        translator,
        files_to_translate,
        source_root_path,
        target_root_path,
        source_lang,
        target_lang,
        TRANSLATION_SUFFIX,
        glossary
    )

    # Получаем финальное использование и рассчитываем реальную стоимость
    try:
        end_usage = translator.get_usage()
        end_chars = end_usage.character.count if end_usage.character.count else 0
        chars_used = end_chars - start_chars

        if chars_used > 0:
            PRICE_PER_MILLION = 20.00  # EUR
            actual_cost = (chars_used / 1_000_000) * PRICE_PER_MILLION

            print(f"\n💰 РЕАЛЬНАЯ СТОИМОСТЬ ПЕРЕВОДА:")
            print(f"  • Использовано символов: {chars_used:,}")
            print(f"  • Стоимость: €{actual_cost:.2f}")
            if estimated_cost > 0:
                accuracy = (actual_cost / estimated_cost) * 100
                print(f"  • Точность оценки: {accuracy:.1f}%")
    except Exception as e:
        print(f"\nНе удалось получить данные о реальной стоимости: {e}")

    # Итоговый отчет
    print("-" * 30)
    print("Итоговый отчет:")
    found_count = len(files_to_translate)
    processed_total = success + warnings + skipped_suffix + skipped_exists + errors

    print(f"Всего найдено файлов ({', '.join(SUPPORTED_EXTENSIONS)}): {found_count}")
    print(f"Всего обработано (включая пропущенные): {processed_total}")
    print(f"  Успешно переведено и очищено (если DOCX): {success}")
    print(f"  Успешно переведено, но с ошибкой пост-обработки (DOCX): {warnings}")
    print(f"  Пропущено (уже имели суффикс '{actual_suffix}'): {skipped_suffix}")
    print(f"  Пропущено (переведенный файл уже существует): {skipped_exists}")
    print(f"  Не удалось перевести (ошибки API/доступа/файла): {errors}")

    # Выводим детали ошибок и предупреждений пост-обработки
    if errors > 0 or warnings > 0:
        print("\nДетали ошибок и предупреждений:")
        error_details.sort(key=lambda x: 0 if x.get('status') == 'error' else 1)

        for i, err_info in enumerate(error_details):
            input_file_path_str = str(err_info.get('file', 'Неизвестный файл'))
            try:
                file_display_name = Path(input_file_path_str).relative_to(source_root_path)
            except (ValueError, TypeError):
                file_display_name = input_file_path_str

            status = err_info.get('status', 'error')
            prefix = "Ошибка"
            if status == "warning":
                prefix = "Предупреждение"

            error_message_full = str(err_info.get('error', 'Нет деталей'))
            error_message_short = error_message_full.splitlines()[0]

            print(f"  {i + 1}. Файл: {file_display_name}")
            print(f"     {prefix}: {error_message_short}")

    print(f"\nПереведенные файлы сохранены в: {target_root_path}")
    print("-" * 30)


# --- Main Execution ---

if __name__ == "__main__":
    start_total_time = time.time()

    # Парсим аргументы командной строки
    args = parse_arguments()

    # Проверяем API ключ
    check_api_key_placeholder()
    translator = initialize_translator(DEEPL_API_KEY)

    # Проверяем, запущен ли скрипт в автоматическом режиме
    if args.no_interactive and args.input:
        # === АВТОМАТИЧЕСКИЙ РЕЖИМ ===
        input_path = Path(args.input)

        if not input_path.exists():
            print(f"❌ Ошибка: Файл не найден: {input_path}")
            sys.exit(1)

        if not input_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            print(f"❌ Ошибка: Неподдерживаемый формат файла: {input_path.suffix}")
            print(f"Поддерживаемые форматы: {', '.join(SUPPORTED_EXTENSIONS)}")
            sys.exit(1)

        # Определяем выходной файл
        if args.output:
            output_path = Path(args.output)
        else:
            target_suffix = f"_to_{args.target.lower().replace('-', '_')}"
            output_path = input_path.with_name(f"{input_path.stem}{target_suffix}{input_path.suffix}")

        print(f"\n=== АВТОМАТИЧЕСКИЙ РЕЖИМ ===")
        print(f"Входной файл: {input_path}")
        print(f"Выходной файл: {output_path}")
        print(f"Направление: {args.source} -> {args.target}")

        # Получаем глоссарий если применимо
        glossary = get_or_create_glossary(translator, args.source, args.target)

        # Выполняем перевод
        result = translate_single_document(
            input_path=input_path,
            output_path=output_path,
            translator=translator,
            target_lang=args.target,
            source_lang=args.source,
            file_index=1,
            total_files=1,
            source_root_path=input_path.parent,
            glossary=glossary
        )

        # Выводим результат и возвращаем код выхода
        if result['status'] == 'success':
            print("✅ Файл успешно переведен!")
            sys.exit(0)
        elif result['status'] == 'success_with_postprocessing_error':
            print("⚠️ Файл переведен, но возникла ошибка при пост-обработке.")
            sys.exit(0)  # Все равно считаем успехом для workflow
        else:
            print("❌ Ошибка при переводе файла.")
            if 'error' in result:
                print(f"Детали: {result['error']}")
            sys.exit(1)

    else:
        # === ИНТЕРАКТИВНЫЙ РЕЖИМ ===
        # Главное меню
        while True:
            print("\n" + "=" * 50)
            print("DeepL Document Translator - Последовательный режим")
            print("=" * 50)
            print("Выберите режим работы:")
            print("  1. Перевести один файл")
            print("  2. Перевести все файлы в папке (рекурсивно)")
            print("  3 Translate_politics. Найти 10 самых больших файлов")
            print("  4. Выход")

            mode_choice = input("\nВаш выбор (1, 2, 3 Translate_politics или 4): ").strip()

            if mode_choice == '1':
                handle_single_file(translator)
            elif mode_choice == '2':
                handle_folder_translation(translator)
            elif mode_choice == '3 Translate_politics':
                find_largest_files()
            elif mode_choice == '4':
                print("\nВыход из программы.")
                break
            else:
                print("Неверный выбор. Пожалуйста, введите 1, 2, 3 Translate_politics или 4.")

        end_total_time = time.time()
        print(f"\nОбщее время выполнения: {end_total_time - start_total_time:.2f} секунд.")