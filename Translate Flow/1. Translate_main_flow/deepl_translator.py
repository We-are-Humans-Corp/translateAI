import deepl
import os
import sys
import time
import traceback
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
import re  # Для регулярных выражений
from dotenv import load_dotenv

# Импортируем словарь глоссария из локального файла
try:
    from deepl_glossary_python import glossary_entries
except ImportError:
    print("Предупреждение: Не удалось импортировать глоссарий из deepl_glossary_python.py")
    print("Перевод будет выполнен без использования глоссария.")
    glossary_entries = None

# --- Проверка и импорт python-docx ---
try:
    from docx import Document
    from docx.shared import Pt  # Для возможной работы со стилями, если понадобится
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
except ImportError:
    print("Ошибка: Необходима библиотека python-docx.")
    print("Пожалуйста, установите ее: pip install python-docx")
    sys.exit(1)
# --- Конец добавления ---

# --- Загрузка переменных окружения ---
# Определяем путь к .env файлу (теперь нужно подняться на 3 Translate_politics уровня вверх)
BASE_DIR = Path(__file__).resolve().parent.parent.parent
load_dotenv(BASE_DIR / '.env')

# --- Configuration ---
# Получаем ключ из переменных окружения
DEEPL_API_KEY = os.getenv('DEEPL_API_KEY', '')

SUPPORTED_EXTENSIONS = ['.docx', '.pdf']
TRANSLATION_SUFFIX = "_to_{target_lang_code}"
MAX_CONCURRENT_TRANSLATIONS = 2  # Можно настроить

LANGUAGE_OPTIONS = {
    "1": {"name": "Русский -> Английский (US)", "source": "RU", "target": "EN-US"},
    "2": {"name": "Английский -> Русский", "source": "EN", "target": "RU"},
    "3 Translate_politics": {"name": "Немецкий -> Русский", "source": "DE", "target": "RU"},
    # Добавьте другие языки при необходимости
}


# --- Helper Functions ---

def check_api_key_placeholder():
    """Checks if the API key is present in environment variables."""
    if not DEEPL_API_KEY:
        print("Ошибка: Ключ DeepL API не найден в переменных окружения.")
        print("Пожалуйста, создайте файл .env в корне проекта и добавьте:")
        print("DEEPL_API_KEY=your_api_key_here")
        print("\nИли установите переменную окружения DEEPL_API_KEY")
        sys.exit(1)


def get_user_input():
    """Gets source path and translation direction from the user."""
    print("--- Скрипт рекурсивного перевода документов через DeepL API ---")
    print("\nВведите ПОЛНЫЙ путь к КОРНЕВОЙ папке с файлами для перевода")
    print(f"(Поддерживаемые форматы: {', '.join(SUPPORTED_EXTENSIONS)})")
    print("Пример для Mac: /Users/имя/Desktop/МоиДокументы")
    print("Пример для Windows: C:\\Users\\имя\\Documents\\МоиДокументы")

    while True:
        source_root_str = input("-> ").strip()
        source_root_path = Path(source_root_str)
        if source_root_path.is_dir():
            break
        else:
            print(f"Ошибка: Папка не найдена или не является папкой: '{source_root_str}'")
            print("Пожалуйста, введите корректный путь.")

    print("\nВыберите направление перевода:")
    for key, value in LANGUAGE_OPTIONS.items():
        print(f"  {key}: {value['name']}")

    while True:
        choice = input(f"Введите номер ({', '.join(LANGUAGE_OPTIONS.keys())}): ").strip()
        if choice in LANGUAGE_OPTIONS:
            selected_lang = LANGUAGE_OPTIONS[choice]
            print(f"Выбрано: {selected_lang['name']}")
            break
        else:
            print("Ошибка: Неверный выбор. Пожалуйста, введите один из предложенных номеров.")

    return source_root_path, selected_lang['source'], selected_lang['target']


def initialize_translator(api_key):
    """Initializes the DeepL translator and checks usage."""
    print("\nИнициализация переводчика DeepL...")
    try:
        translator = deepl.Translator(api_key)
        print("Проверка лимитов DeepL API...")
        usage = translator.get_usage()

        # Character limit check
        char_count = usage.character.count
        char_limit = usage.character.limit
        limit_reached_char = False
        if char_limit is not None:
            count_str = f"{char_count:,}" if char_count is not None else "N/A"
            limit_str = f"{char_limit:,}"
            remaining_str = ""
            if char_count is not None:
                remaining_str = f"(Осталось: {max(0, char_limit - char_count):,})"
                if char_count is not None and char_limit is not None and char_count >= char_limit: limit_reached_char = True
            print(f"  Лимит символов: Использовано {count_str} из {limit_str} {remaining_str}")
            if limit_reached_char: print("Внимание: Лимит символов для перевода исчерпан или достигнут!")
        elif char_count is not None:
            print(f"  Лимит символов: Использовано {char_count:,} (Лимит не установлен или не известен)")
        else:
            print("  Лимит символов: Использование не отслеживается или не применимо.")

        # Document limit check
        doc_count = None
        doc_limit = None
        limit_reached_doc = False
        if hasattr(usage, 'document') and usage.document and usage.document.valid:
            doc_count = usage.document.count
            doc_limit = usage.document.limit
            if doc_limit is not None:
                count_str = f"{doc_count:,}" if doc_count is not None else "N/A"
                limit_str = f"{doc_limit:,}"
                remaining_str = ""
                if doc_count is not None:
                    remaining_str = f"(Осталось: {max(0, doc_limit - doc_count):,})"
                    if doc_count is not None and doc_limit is not None and doc_count >= doc_limit: limit_reached_doc = True
                print(f"  Лимит документов: Использовано {count_str} из {limit_str} {remaining_str}")
                if limit_reached_doc: print("Внимание: Лимит документов для перевода исчерпан или достигнут!")
            elif doc_count is not None:
                print(f"  Лимит документов: Использовано {doc_count:,} (Лимит не установлен или не известен)")
            else:
                print("  Лимит документов: Использование не отслеживается или не применимо для этого плана.")
        else:
            print("  Лимит документов: Не применимо или не предоставлено API (возможно, DeepL API Free?)")

        # Final check if any limit is exceeded
        if usage.any_limit_exceeded or limit_reached_char or limit_reached_doc:
            error_msg = "Достигнут лимит использования DeepL API."
            if limit_reached_char: error_msg = "Достигнут лимит символов DeepL API."
            if limit_reached_doc: error_msg = "Достигнут лимит документов DeepL API."
            print(f"\nОшибка: {error_msg}")
            print("Перевод невозможен.")
            sys.exit(1)
        else:
            print("Инициализация DeepL успешна. Лимиты в норме.")
        return translator

    except deepl.AuthorizationException:
        print("\nОшибка аутентификации DeepL: Неверный API ключ.")
        sys.exit(1)
    except Exception as e:
        print(f"\nНепредвиденная ошибка при инициализации DeepL: {e}")
        print(traceback.format_exc())
        sys.exit(1)


def get_or_create_glossary(translator, source_lang, target_lang):
    """Получает существующий глоссарий или создает новый."""
    # Если глоссарий не был импортирован, возвращаем None
    if glossary_entries is None:
        return None

    # Проверяем, что направление перевода поддерживается глоссарием
    if source_lang != "RU" or target_lang not in ["EN-US", "EN-GB", "EN"]:
        print("Глоссарий доступен только для перевода с русского на английский.")
        return None

    glossary_name = f"Universal Scientific Terms RU-EN v1"

    try:
        # Пытаемся найти существующий глоссарий
        print("\nПроверка существующих глоссариев...")
        glossaries = translator.list_glossaries()

        for glossary in glossaries:
            if glossary.name == glossary_name and glossary.source_lang == "RU" and glossary.target_lang == "EN":
                print(f"✅ Найден существующий глоссарий: {glossary.name} (ID: {glossary.glossary_id})")
                print(f"   Количество терминов: {glossary.entry_count}")
                return glossary

        # Если глоссарий не найден, создаем новый
        print("Создание нового глоссария научных терминов...")
        glossary = translator.create_glossary(
            name=glossary_name,
            source_lang="RU",
            target_lang="EN",
            entries=glossary_entries
        )
        print(f"✅ Глоссарий успешно создан: {glossary.name} (ID: {glossary.glossary_id})")
        print(f"   Количество терминов: {len(glossary_entries)}")
        return glossary

    except deepl.DeepLException as e:
        print(f"⚠️ Не удалось создать/получить глоссарий: {e}")
        print("   Перевод будет выполнен без использования глоссария.")
        return None
    except Exception as e:
        print(f"⚠️ Неожиданная ошибка при работе с глоссарием: {e}")
        print("   Перевод будет выполнен без использования глоссария.")
        return None


def find_files_to_translate(source_root_path):
    """Recursively finds files with supported extensions, excluding temporary files."""
    print("\nПоиск файлов для перевода...")
    files_found = []
    for extension in SUPPORTED_EXTENSIONS:
        extension_pattern = "".join([f"[{c.lower()}{c.upper()}]" for c in extension[1:]])
        pattern = f"*.{extension_pattern}"
        # print(f"Поиск файлов с шаблоном: {pattern} в {source_root_path}") # Debug log
        files_found.extend(list(source_root_path.rglob(pattern)))

    processed_paths = set()
    files_to_process = []
    for f in files_found:
        try:
            resolved_f = f.resolve()
            if resolved_f.is_file() and not resolved_f.name.startswith('~$') and resolved_f not in processed_paths:
                files_to_process.append(resolved_f)
                processed_paths.add(resolved_f)
        except Exception as e:
            print(f"Предупреждение: Не удалось обработать путь файла '{f}': {e}. Пропуск.")

    files_to_process.sort()
    print(f"Найдено {len(files_to_process)} файлов с расширениями {SUPPORTED_EXTENSIONS} (исключая временные файлы).")
    return files_to_process


# --- ОБНОВЛЕННАЯ ФУНКЦИЯ ОЧИСТКИ (Версия 3 Translate_politics) ---
def clean_translated_docx(docx_path):
    """
    Открывает переведенный DOCX и исправляет искаженные плейсхолдеры
    к виду <<EqnXXX.eps>>. Обрабатывает случаи с лишней точкой,
    отсутствующими >>, и отсутствующим .eps.
    Перезаписывает файл. Возвращает True при успехе/отсутствии изменений, False при ошибке.
    """
    try:
        if not docx_path.is_file():
            print(f"  -> ОШИБКА Пост-обработки: Файл не найден {docx_path}")
            return False

        document = Document(docx_path)
        changes_made = 0  # Счетчик изменений для логирования

        # УЛУЧШЕННОЕ РЕГУЛЯРНОЕ ВЫРАЖЕНИЕ v3:
        pattern = re.compile(
            r'(<<Eqn(\d+))'  # Группа 1: <<Eqn<цифры>, Группа 2: только цифры
            r'(?!'  # Начало негативного просмотра вперед (убеждаемся, что ДАЛЬШЕ НЕ...)
            r'\.eps>>'  # ...ровно ".eps>>"
            r'([,\s]|$)'  # ...за которым идет запятая, пробел или конец строки/параграфа
            r')'  # Конец негативного просмотра
            r'([\.\w>]+)?'  # Группа 3 Translate_politics (опциональная): Захватываем сам "мусор" - точки, буквы(eps), >.
        )

        # Строка для замены: Восстанавливаем правильный формат, используя Группу 2 (цифры)
        replacement_pattern = r'<<Eqn\2.eps>>'

        # Вспомогательная функция для обработки параграфа
        def process_paragraph(para: Paragraph):
            nonlocal changes_made
            # Быстрая проверка для оптимизации - ищем хотя бы начало плейсхолдера
            if '<<Eqn' not in para.text:
                return False

            original_text = para.text
            # Применяем замену ко всему тексту параграфа
            new_text = pattern.sub(replacement_pattern, original_text)

            if new_text != original_text:
                # Считаем количество реальных замен (хотя бы примерно)
                # Это не идеально точно, но дает представление
                num_replacements = len(pattern.findall(original_text))
                changes_made += num_replacements

                # Простой способ обновления: очистить параграф и вставить новый текст
                para.clear()  # Очищает все runs и текст
                para.add_run(new_text)  # Добавляет новый текст одним run (может сбросить форматирование)

                # Debug Log (можно раскомментировать для отладки)
                # print(f"    Debug: Replaced in para. Orig: '{original_text}'. New: '{new_text}'")
                return True
            return False

        # --- Основная логика обхода документа ---
        # Итерация по параграфам в основном тексте
        for para in document.paragraphs: process_paragraph(para)

        # Итерация по таблицам
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs: process_paragraph(para)

        # Итерация по колонтитулам (если нужно)
        for section in document.sections:
            # Headers
            for para in section.header.paragraphs: process_paragraph(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs: process_paragraph(para)
            # Footers
            for para in section.footer.paragraphs: process_paragraph(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs: process_paragraph(para)
        # --- Конец обхода ---

        if changes_made > 0:
            print(f"  -> Пост-обработка V3: Исправлено ~{changes_made} искаженных плейсхолдеров в {docx_path.name}")
            document.save(docx_path)  # Перезаписываем файл
        else:
            print(
                f"  -> Пост-обработка V3: Искаженные плейсхолдеры (требующие исправления) не найдены в {docx_path.name}")

        return True

    except FileNotFoundError:
        print(f"  -> ОШИБКА Пост-обработки: Файл не найден {docx_path}")
        return False
    except Exception as e:
        # Используем f-string для форматирования
        print(f"  -> КРИТИЧЕСКАЯ ОШИБКА Пост-обработки: Не удалось обработать {docx_path.name}: {e}")
        print(traceback.format_exc())  # Для отладки
        return False


# --- КОНЕЦ ОБНОВЛЕННОЙ ФУНКЦИИ ---

def fix_chemical_formulas_in_docx(docx_path):
    """
    Исправляет химические формулы в переведенных DOCX файлах.
    Удаляет скобки из чисел в нижнем индексе (subscript).
    Например: H(2)O -> H2O (где 2 остается в subscript)

    Args:
        docx_path: Path объект к DOCX файлу

    Returns:
        bool: True если обработка прошла успешно, False в случае ошибки
    """
    try:
        if not docx_path.is_file():
            print(f"  -> ОШИБКА: Файл не найден {docx_path}")
            return False

        document = Document(docx_path)
        total_fixes = 0

        def process_runs_in_paragraph(paragraph):
            """Обрабатывает все run в параграфе для исправления химических формул"""
            nonlocal total_fixes
            fixes_in_para = 0

            # Создаем новый список runs для замены
            new_runs = []

            for run in paragraph.runs:
                if run.font.subscript:
                    # Это subscript текст
                    original_text = run.text

                    # Удаляем скобки из текста
                    # Паттерн 1: (число) -> число
                    new_text = re.sub(r'^\((\d+(?:\.\d+)?)\)$', r'\1', original_text)

                    # Паттерн 2: число) -> число
                    new_text = re.sub(r'^(\d+(?:\.\d+)?)\)$', r'\1', new_text)

                    # Паттерн 3 Translate_politics: (число -> число
                    new_text = re.sub(r'^\((\d+(?:\.\d+)?)$', r'\1', new_text)

                    # Паттерн 4: одиночная скобка
                    if new_text in ['(', ')']:
                        new_text = ''

                    if new_text != original_text:
                        fixes_in_para += 1

                        # Создаем новый run с исправленным текстом
                        new_run = paragraph.add_run(new_text)
                        # Копируем форматирование
                        new_run.font.subscript = True
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.bold:
                            new_run.font.bold = run.font.bold
                        if run.font.italic:
                            new_run.font.italic = run.font.italic
                        new_runs.append(new_run)
                    else:
                        new_runs.append(run)
                else:
                    # Обычный текст - проверяем на наличие формул
                    text = run.text

                    # Проверяем, есть ли химические формулы с числами не в subscript
                    # Например: Ti(49) -> Ti49 (где 49 должно быть в subscript)
                    pattern = re.compile(r'([A-Z][a-z]?)\((\d+(?:\.\d+)?)\)')

                    if pattern.search(text):
                        # Есть формула для исправления
                        parts = []
                        last_end = 0

                        for match in pattern.finditer(text):
                            # Добавляем текст до формулы
                            if match.start() > last_end:
                                part_run = paragraph.add_run(text[last_end:match.start()])
                                # Копируем форматирование
                                if run.font.size:
                                    part_run.font.size = run.font.size
                                if run.font.name:
                                    part_run.font.name = run.font.name
                                if run.font.bold:
                                    part_run.font.bold = run.font.bold
                                if run.font.italic:
                                    part_run.font.italic = run.font.italic
                                parts.append(part_run)

                            # Добавляем элемент
                            element = match.group(1)
                            element_run = paragraph.add_run(element)
                            # Копируем форматирование
                            if run.font.size:
                                element_run.font.size = run.font.size
                            if run.font.name:
                                element_run.font.name = run.font.name
                            if run.font.bold:
                                element_run.font.bold = run.font.bold
                            if run.font.italic:
                                element_run.font.italic = run.font.italic
                            parts.append(element_run)

                            # Добавляем число в subscript
                            number = match.group(2)
                            number_run = paragraph.add_run(number)
                            number_run.font.subscript = True
                            if run.font.size:
                                number_run.font.size = run.font.size
                            if run.font.name:
                                number_run.font.name = run.font.name
                            parts.append(number_run)

                            last_end = match.end()
                            fixes_in_para += 1

                        # Добавляем оставшийся текст
                        if last_end < len(text):
                            part_run = paragraph.add_run(text[last_end:])
                            # Копируем форматирование
                            if run.font.size:
                                part_run.font.size = run.font.size
                            if run.font.name:
                                part_run.font.name = run.font.name
                            if run.font.bold:
                                part_run.font.bold = run.font.bold
                            if run.font.italic:
                                part_run.font.italic = run.font.italic
                            parts.append(part_run)

                        new_runs.extend(parts)
                    else:
                        new_runs.append(run)

            if fixes_in_para > 0:
                # Очищаем параграф и добавляем новые runs
                paragraph.clear()
                paragraph._p.extend([run._element for run in new_runs])
                total_fixes += fixes_in_para
                return True

            return False

        # Обрабатываем все части документа
        # Параграфы
        for para in document.paragraphs:
            process_runs_in_paragraph(para)

        # Таблицы
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_runs_in_paragraph(para)

        # Колонтитулы
        for section in document.sections:
            # Headers
            for para in section.header.paragraphs:
                process_runs_in_paragraph(para)
            # Footers
            for para in section.footer.paragraphs:
                process_runs_in_paragraph(para)

        # Сохраняем результат
        if total_fixes > 0:
            document.save(docx_path)
            print(f"  -> Исправление химических формул: исправлено {total_fixes} формул")
            return True
        else:
            return True  # Изменений не требуется, но это не ошибка

    except Exception as e:
        print(f"  -> ОШИБКА при исправлении химических формул: {e}")
        return False


def translate_single_document(input_path, output_path, translator, target_lang, source_lang, file_index, total_files,
                              source_root_path, glossary=None, sequential_mode=False):
    """Translates a single document and handles errors. Includes enhanced post-processing for DOCX."""
    # Конвертируем Path объекты в строки для DeepL API и логгирования
    input_path_str = str(input_path)
    output_path_str = str(output_path)

    try:
        relative_path_for_display = input_path.relative_to(source_root_path)
    except ValueError:
        relative_path_for_display = input_path
    except Exception as e:
        relative_path_for_display = f"Ошибка при определении пути: {input_path_str}"  # Используем str
        print(f"Предупреждение: Не удалось определить относительный путь для отображения: {e}")

    # Используем имена файлов из Path объектов для большей точности
    print(f"\n[{file_index}/{total_files}] Обработка файла: {relative_path_for_display}")
    print(f"  Перевод '{input_path.name}' -> '{output_path.name}'...")
    print(f"     (Языки: {source_lang} -> {target_lang})")

    post_processing_error_occurred = False  # Флаг для отслеживания ошибок пост-обработки

    try:
        # Используем Path объект для создания директории
        output_path.parent.mkdir(parents=True, exist_ok=True)

        start_time = time.time()

        # Добавляем логирование использования глоссария
        if glossary:
            print(f"     (Используется глоссарий: {glossary.name})")

        # Передаем глоссарий в функцию перевода, если он доступен
        translation_kwargs = {
            "input_path": input_path_str,  # DeepL требует строку
            "output_path": output_path_str,  # DeepL требует строку
            "target_lang": target_lang,
            "source_lang": source_lang
        }

        if glossary:
            translation_kwargs["glossary"] = glossary

        translator.translate_document_from_filepath(**translation_kwargs)

        end_time = time.time()
        print(f"  -> УСПЕШНО переведен за {end_time - start_time:.2f} сек.")

        # Добавляем паузу между запросами в последовательном режиме
        if sequential_mode:
            print(f"  -> Пауза 1 секунда перед следующим запросом...")
            time.sleep(1)

        # --- ИНТЕГРАЦИЯ ПОСТ-ОБРАБОТКИ ---
        if output_path.suffix.lower() == '.docx':
            print(f"  -> Запуск пост-обработки V3 для {output_path.name}...")
            # Передаем Path объект в функцию очистки
            cleaning_successful = clean_translated_docx(output_path)
            if not cleaning_successful:
                # Логируем предупреждение и устанавливаем флаг
                print(f"  -> ПРЕДУПРЕЖДЕНИЕ: Пост-обработка файла {output_path.name} завершилась с ошибкой.")
                post_processing_error_occurred = True  # Отмечаем ошибку пост-обработки

            # Добавляем исправление химических формул
            print(f"  -> Проверка и исправление химических формул...")
            formulas_fixed = fix_chemical_formulas_in_docx(output_path)
            if not formulas_fixed:
                print(f"  -> ПРЕДУПРЕЖДЕНИЕ: Исправление химических формул завершилось с ошибкой.")
                post_processing_error_occurred = True
        # --- КОНЕЦ ИНТЕГРАЦИИ ---

        # Возвращаем статус и пути. Добавляем информацию об ошибке пост-обработки.
        return {
            "status": "success" if not post_processing_error_occurred else "success_with_postprocessing_error",
            "input": input_path_str,
            "output": output_path_str,
            "post_processing_error": post_processing_error_occurred  # Явно указываем на проблему
        }

    # Обработка ошибок DeepL API и файловой системы
    except deepl.DocumentTranslationException as e:
        error_msg = f"Ошибка перевода документа DeepL: {e}"
        print(f"  -> ОШИБКА: {error_msg}")
        # ... (дополнительные детали ошибки DeepL, если нужны) ...
        return {"status": "error", "input": input_path_str, "error": error_msg}
    except deepl.QuotaExceededException:
        error_msg = "Превышена квота DeepL API."
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg, "quota_exceeded": True}
    except deepl.TooManyRequestsException:
        error_msg = "Слишком много запросов к DeepL API."
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg, "rate_limited": True}
    except deepl.ConnectionException as e:
        error_msg = f"Ошибка сети при обращении к DeepL: {e}"
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg}
    except deepl.DeepLException as e:  # Общая ошибка DeepL
        error_msg = f"Общая ошибка DeepL API: {e}"
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg}
    except FileNotFoundError:
        # Скорее всего, не найден ИСХОДНЫЙ файл
        error_msg = f"Исходный файл не найден: {input_path_str}"
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg}
    except PermissionError:
        error_msg = f"Нет прав на чтение/запись файла ({input_path_str} -> {output_path_str})."
        print(f"  -> ОШИБКА: {error_msg}")
        return {"status": "error", "input": input_path_str, "error": error_msg}
    except Exception as e:
        # Перехват других непредвиденных ошибок (включая возможные ошибки при вызове clean_translated_docx)
        error_msg = f"Непредвиденная ошибка при обработке файла '{input_path_str}': {e}"
        print(f"  -> КРИТИЧЕСКАЯ ОШИБКА: {error_msg}")
        print(traceback.format_exc())
        return {"status": "error", "input": input_path_str, "output": output_path_str,
                "error": f"{error_msg}\n{traceback.format_exc()}"}


def process_translations(translator, file_paths, source_root_path, target_root_path, source_lang, target_lang,
                         suffix_pattern, glossary=None, sequential_mode=False):
    """Manages the translation process, using parallelism or sequential processing."""
    success_count = 0
    success_with_warnings = 0  # Отдельный счетчик для успехов с ошибками пост-обработки
    skipped_suffix_count = 0
    skipped_exists_count = 0
    error_count = 0
    errors_list = []  # Список деталей ошибок
    total_files = len(file_paths)
    stop_processing = False  # Флаг для остановки при критических ошибках API

    actual_suffix = suffix_pattern.format(target_lang_code=target_lang.lower().replace("-", "_"))

    # Устанавливаем количество потоков в зависимости от режима
    max_workers = 1 if sequential_mode else MAX_CONCURRENT_TRANSLATIONS
    processing_mode = "ПОСЛЕДОВАТЕЛЬНОЙ" if sequential_mode else f"параллельной (до {MAX_CONCURRENT_TRANSLATIONS} потоков)"

    print(f"\nЗапуск {processing_mode} обработки {total_files} файлов...")
    print("-" * 30)

    executor = ThreadPoolExecutor(max_workers=max_workers)
    future_to_path_obj = {}  # Храним Path объекты
    files_submitted = 0

    for i, input_path_obj in enumerate(file_paths):  # Используем Path объект
        if stop_processing:
            print(
                f"\n[{i + 1}/{total_files}] Пропуск отправки файла {input_path_obj.name} из-за критической ошибки API.")
            continue

        file_index = i + 1
        try:
            relative_path_for_target = input_path_obj.relative_to(source_root_path)
            output_path_obj = target_root_path / relative_path_for_target.with_name(
                f"{input_path_obj.stem}{actual_suffix}{input_path_obj.suffix}")
        except Exception as e:
            print(
                f"\n[{file_index}/{total_files}] Ошибка при расчете выходного пути для {input_path_obj.name}: {e}. Пропуск файла.")
            error_count += 1
            # Добавляем детали ошибки
            errors_list.append(
                {"status": "error", "file": str(input_path_obj), "error": f"Ошибка при расчете выходного пути: {e}"})
            continue

        # --- Skip Checks ---
        if input_path_obj.stem.endswith(actual_suffix):
            try:
                display_path = input_path_obj.relative_to(source_root_path)
            except ValueError:
                display_path = input_path_obj.name
            print(f"\n[{file_index}/{total_files}] Пропуск (файл уже имеет суффикс '{actual_suffix}'): {display_path}")
            skipped_suffix_count += 1
            continue

        if output_path_obj.exists():
            try:
                display_path = output_path_obj.relative_to(target_root_path)
            except ValueError:
                display_path = output_path_obj.name
            print(f"\n[{file_index}/{total_files}] Пропуск (переведенный файл уже существует): {display_path}")
            skipped_exists_count += 1
            continue
        # --- End Skip Checks ---

        # Submit the translation task
        future = executor.submit(
            translate_single_document,
            input_path_obj,  # Передаем Path
            output_path_obj,  # Передаем Path
            translator,
            target_lang,
            source_lang,
            file_index,
            total_files,
            source_root_path,  # Path
            glossary,  # Передаем глоссарий
            sequential_mode  # Передаем режим обработки
        )
        future_to_path_obj[future] = input_path_obj  # Храним Path объект
        files_submitted += 1

    print(f"\nОтправлено {files_submitted} заданий на перевод. Ожидание завершения...")

    # --- Сбор результатов ---
    try:
        for future in as_completed(future_to_path_obj):
            input_path_obj = future_to_path_obj[future]
            input_filename_str = str(input_path_obj)  # Строка для логов
            try:
                result = future.result()  # Получаем результат из потока

                if result['status'] == "success":
                    success_count += 1
                elif result['status'] == "success_with_postprocessing_error":
                    # Успешный перевод, но проблема с очисткой
                    success_with_warnings += 1
                    # Добавляем информацию об ошибке пост-обработки в общий список ошибок
                    errors_list.append({
                        "status": "warning",  # Используем статус warning
                        "file": result.get('input', input_filename_str),
                        "output": result.get('output'),
                        "error": f"Ошибка пост-обработки файла {Path(result.get('output', '')).name}"
                    })
                elif result['status'] == "error":
                    # Ошибка перевода или другая критическая ошибка
                    error_count += 1
                    # Убедимся, что ключ 'file' есть в словаре ошибки для отчета
                    if 'file' not in result: result['file'] = result.get('input', input_filename_str)
                    errors_list.append(result)
                    # Проверяем на критические ошибки API для остановки
                    if result.get("quota_exceeded") or result.get("rate_limited"):
                        if not stop_processing:
                            print("\n*** Обнаружено превышение квоты или лимита запросов DeepL. ***")
                            print("*** Новые задания не будут отправляться. Дождитесь завершения текущих. ***")
                        stop_processing = True
                else:
                    # Непредвиденный статус
                    print(
                        f"Предупреждение: Неизвестный статус результата '{result.get('status')}' для файла {input_filename_str}")
                    error_count += 1  # Считаем как ошибку
                    errors_list.append({"status": "error", "file": input_filename_str,
                                        "error": f"Неизвестный статус результата: {result.get('status')}"})


            except Exception as exc:
                # Ошибка получения результата из самого future (редко)
                error_count += 1
                error_msg = f"Критическая ошибка при получении результата для файла '{input_filename_str}': {exc}"
                print(f"  -> КРИТИЧЕСКАЯ ОШИБКА ПОТОКА: {error_msg}")
                print(traceback.format_exc())
                errors_list.append(
                    {"status": "error", "file": input_filename_str, "error": f"{error_msg}\n{traceback.format_exc()}"})
    finally:
        # Гарантированно закрываем пул потоков
        executor.shutdown(wait=True)

    print("-" * 30)
    print("Обработка завершена.")
    # Возвращаем все счетчики
    return success_count, success_with_warnings, skipped_suffix_count, skipped_exists_count, error_count, errors_list


# --- НОВЫЕ ФУНКЦИИ ДЛЯ РЕЖИМА ОДНОГО ФАЙЛА ---

def handle_single_file(translator):
    """Обработка перевода одного файла"""
    print("\n--- Режим перевода одного файла ---")

    # 1. Запрашиваем путь к файлу
    while True:
        source_str = input("\nВведите ПОЛНЫЙ путь к файлу для перевода: ").strip().strip('"\'')
        source_path = Path(source_str)
        if source_path.is_file() and source_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            break
        else:
            print(f"Ошибка: Файл не найден или не поддерживается.")
            print(f"Поддерживаемые форматы: {', '.join(SUPPORTED_EXTENSIONS)}")

    # 2. Запрашиваем направление перевода
    print("\nВыберите направление перевода:")
    for key, value in LANGUAGE_OPTIONS.items():
        print(f"  {key}: {value['name']}")

    while True:
        choice = input(f"Введите номер ({', '.join(LANGUAGE_OPTIONS.keys())}): ").strip()
        if choice in LANGUAGE_OPTIONS:
            selected_lang = LANGUAGE_OPTIONS[choice]
            break
        else:
            print("Ошибка: Неверный выбор.")

    source_lang = selected_lang['source']
    target_lang = selected_lang['target']

    # 3 Translate_politics. Получаем или создаем глоссарий
    glossary = get_or_create_glossary(translator, source_lang, target_lang)

    # 4. Определяем имя выходного файла
    target_lang_code_for_suffix = target_lang.lower().replace("-", "_")
    actual_suffix = TRANSLATION_SUFFIX.format(target_lang_code=target_lang_code_for_suffix)
    output_path = source_path.with_name(f"{source_path.stem}{actual_suffix}{source_path.suffix}")

    print("-" * 50)
    print(f"Исходный файл: {source_path}")
    print(f"Целевой файл: {output_path}")
    print(f"Направление: {source_lang} -> {target_lang}")
    if glossary:
        print(f"Глоссарий: {glossary.name}")
    print("-" * 50)

    # 5. Проверяем существование файла
    if output_path.exists():
        overwrite = input(f"\nФайл '{output_path.name}' уже существует. Перезаписать? (y/n): ").strip().lower()
        if overwrite != 'y':
            print("Перевод отменен.")
            return

    # 6. Выполняем перевод
    result = translate_single_document(
        input_path=source_path,
        output_path=output_path,
        translator=translator,
        target_lang=target_lang,
        source_lang=source_lang,
        file_index=1,
        total_files=1,
        source_root_path=source_path.parent,
        glossary=glossary,
        sequential_mode=True
    )

    # 7. Выводим результат
    print("\n" + "=" * 50)
    if result['status'] == "success":
        print("✅ Файл успешно переведен!")
        print(f"Сохранен как: {output_path}")
    elif result['status'] == "success_with_postprocessing_error":
        print("⚠️ Файл переведен, но возникла ошибка при пост-обработке.")
        print(f"Сохранен как: {output_path}")
    else:
        print("❌ Ошибка при переводе файла.")
        if 'error' in result:
            print(f"Детали: {result['error']}")
    print("=" * 50)


def handle_folder_with_menu(translator):
    """Обработка папки с файлами - оригинальная логика"""
    source_root_path, source_lang, target_lang = get_user_input()

    target_lang_code_for_suffix = target_lang.lower().replace("-", "_")
    actual_suffix = TRANSLATION_SUFFIX.format(target_lang_code=target_lang_code_for_suffix)
    target_root_path = source_root_path.with_name(f"{source_root_path.name}{actual_suffix}")

    print("-" * 30)
    print(f"Исходная корневая папка: {source_root_path}")
    print(f"Целевая корневая папка для переводов: {target_root_path}")
    print(f"Направление перевода: {source_lang} -> {target_lang}")
    print(f"Суффикс переведенных файлов: {actual_suffix}")
    print("-" * 30)

    print(f"\nПроверка/создание папки для переводов: {target_root_path}")
    try:
        target_root_path.mkdir(parents=True, exist_ok=True)
        print("Папка для переводов готова.")
    except PermissionError:
        print(f"Ошибка: Нет прав на создание папки '{target_root_path}'.")
        return
    except Exception as e:
        print(f"Ошибка при создании папки для переводов: {e}")
        print(traceback.format_exc())
        return

    files_to_translate = find_files_to_translate(source_root_path)

    if not files_to_translate:
        print("\nНе найдено подходящих файлов для перевода в указанной папке и ее подпапках.")
        return

    # Спрашиваем пользователя о режиме обработки
    print("\nВыберите режим обработки файлов:")
    print("  1: Параллельная обработка (быстрее, до 2 файлов одновременно)")
    print("  2: Последовательная обработка (медленнее, но безопаснее для API)")

    while True:
        mode_choice = input("Введите номер (1 или 2) [по умолчанию 1]: ").strip()
        if mode_choice == "" or mode_choice == "1":
            sequential_mode = False
            print("Выбрана параллельная обработка.")
            break
        elif mode_choice == "2":
            sequential_mode = True
            print("Выбрана последовательная обработка.")
            break
        else:
            print("Ошибка: Неверный выбор. Пожалуйста, введите 1 или 2.")

    # Получаем или создаем глоссарий (если применимо)
    glossary = get_or_create_glossary(translator, source_lang, target_lang)

    # Запуск процесса перевода и очистки
    success, warnings, skipped_suffix, skipped_exists, errors, error_details = process_translations(
        translator,
        files_to_translate,
        source_root_path,
        target_root_path,
        source_lang,
        target_lang,
        TRANSLATION_SUFFIX,
        glossary,
        sequential_mode
    )

    # Итоговый отчет
    print("-" * 30)
    print("Итоговый отчет:")
    found_count = len(files_to_translate)
    processed_total = success + warnings + skipped_suffix + skipped_exists + errors

    print(f"Всего найдено файлов ({', '.join(SUPPORTED_EXTENSIONS)}): {found_count}")
    print(f"Всего обработано (включая пропущенные): {processed_total}")
    print(f"  Успешно переведено и очищено (если DOCX): {success}")
    print(f"  Успешно переведено, но с ошибкой пост-обработки (DOCX): {warnings}")
    print(f"  Пропущено (уже имели суффикс '{actual_suffix}'): {skipped_suffix}")
    print(f"  Пропущено (переведенный файл уже существует): {skipped_exists}")
    print(f"  Не удалось перевести (ошибки API/доступа/файла): {errors}")

    # Выводим детали ошибок и предупреждений пост-обработки
    if errors > 0 or warnings > 0:
        print("\nДетали ошибок и предупреждений:")
        error_details.sort(key=lambda x: 0 if x.get('status') == 'error' else 1)

        for i, err_info in enumerate(error_details):
            input_file_path_str = str(err_info.get('file', 'Неизвестный файл'))
            try:
                file_display_name = Path(input_file_path_str).relative_to(source_root_path)
            except (ValueError, TypeError):
                file_display_name = input_file_path_str

            status = err_info.get('status', 'error')
            prefix = "Ошибка"
            if status == "warning":
                prefix = "Предупреждение"

            error_message_full = str(err_info.get('error', 'Нет деталей'))
            error_message_short = error_message_full.splitlines()[0]

            print(f"  {i + 1}. Файл: {file_display_name}")
            print(f"     {prefix}: {error_message_short}")

    print(f"\nПереведенные файлы сохранены в: {target_root_path}")
    print("-" * 30)


# --- Main Execution ---

if __name__ == "__main__":
    start_total_time = time.time()

    check_api_key_placeholder()
    translator = initialize_translator(DEEPL_API_KEY)

    # Главное меню
    while True:
        print("\n" + "=" * 50)
        print("DeepL Document Translator")
        print("=" * 50)
        print("Выберите режим работы:")
        print("  1. Перевести один файл")
        print("  2. Перевести все файлы в папке (рекурсивно)")
        print("  3 Translate_politics. Выход")

        mode_choice = input("\nВаш выбор (1, 2 или 3 Translate_politics): ").strip()

        if mode_choice == '1':
            handle_single_file(translator)
        elif mode_choice == '2':
            handle_folder_with_menu(translator)
        elif mode_choice == '3 Translate_politics':
            print("\nВыход из программы.")
            break
        else:
            print("Неверный выбор. Пожалуйста, введите 1, 2 или 3 Translate_politics.")

    end_total_time = time.time()
    print(f"\nОбщее время выполнения: {end_total_time - start_total_time:.2f} секунд.")
    print("Спасибо за использование DeepL Document Translator!")