#!/usr/bin/env python3

"""
EQN RESTORATION BATCH - VERSION V4 (ОКОНЧАТЕЛЬНАЯ ИСПРАВЛЕННАЯ)

Полностью исправленная версия после анализа реальных результатов:

ИСПРАВЛЕННЫЕ КРИТИЧЕСКИЕ ПРОБЛЕМЫ:
1. Правильная обработка .eps в плейсхолдерах
2. Исправление логики замены без создания артефактов  
3. Точное сохранение формата оригинальных плейсхолдеров
4. Предотвращение дублирования символов
5. Корректная обработка запятых
"""

import sys
from pathlib import Path
import re
from datetime import datetime
import json

try:
    from docx import Document
except ImportError:
    print("Ошибка: Необходима библиотека python-docx.")
    print("Установите: pip install python-docx")
    sys.exit(1)

class PlaceholderRestorer:
    """Окончательно исправленный класс для восстановления последовательности плейсхолдеров"""

    def __init__(self):
        self.stats = {
            'placeholders_replaced': 0,
            'placeholders_found_in_translation': 0,
            'placeholders_found_in_original': 0,
            'damaged_placeholders_fixed': 0,
            'damaged_placeholders_details': []
        }

        # Паттерн для поиска ВСЕХ плейсхолдеров (правильных и поврежденных)
        self.pattern_all_placeholders = re.compile(
            r'(?:<<|<\s*<|<|^|\s)'           # начало 
            r'(Eqn\d+(?:\.eps)?)'             # основная часть (группа 1)
            r'(?:>>>|>>|>\s*>|>|(?=\s)|(?=$)|(?=[,.]))'  # конец
            r'(?:[,.>\s]*)?',                  # возможные артефакты после
            re.IGNORECASE | re.MULTILINE
        )

    def extract_placeholders_list(self, doc_path):
        """
        Извлекает список всех плейсхолдеров из документа в порядке появления
        КРИТИЧЕСКИ ВАЖНО: сохраняет точный формат из оригинала
        """
        placeholders = []

        try:
            document = Document(doc_path)
        except Exception as e:
            print(f"⚠️ Не удалось открыть {doc_path.name}: {e}")
            return []

        def extract_from_text(text):
            """Извлекает плейсхолдеры из текста с сохранением точного формата"""
            if not text:
                return []

            # Ищем все варианты плейсхолдеров
            found_placeholders = []

            # Сначала ищем правильно оформленные
            correct_pattern = re.compile(r'<<Eqn\d+(?:\.eps)?>>(?:,)?', re.IGNORECASE)
            for match in correct_pattern.finditer(text):
                found_placeholders.append((match.start(), match.end(), match.group()))

            # Затем ищем поврежденные в местах, где нет правильных
            damaged_patterns = [
                r'(?<!<)Eqn\d+(?:\.eps)?>>(?:,)?',  # без <<
                r'<<Eqn\d+(?:\.eps)?(?!>>)',        # без >>  
                r'<\s+<Eqn\d+(?:\.eps)?>\s*>',    # с пробелами
                r'<<Eqn\d+(?:\.eps)?>>>+(?:,)?',    # лишние >
            ]

            for pattern in damaged_patterns:
                for match in re.finditer(pattern, text, re.IGNORECASE):
                    start, end = match.start(), match.end()
                    # Проверяем, не пересекается ли с уже найденными
                    overlap = False
                    for existing_start, existing_end, _ in found_placeholders:
                        if not (end <= existing_start or start >= existing_end):
                            overlap = True
                            break

                    if not overlap:
                        # Исправляем поврежденный плейсхолдер
                        original = match.group()
                        eqn_match = re.search(r'Eqn(\d+)', original, re.IGNORECASE)
                        if eqn_match:
                            num = eqn_match.group(1)
                            has_eps = '.eps' in original.lower()
                            has_comma = original.endswith(',')

                            if has_eps:
                                fixed = f'<<Eqn{num}.eps>>'
                            else:
                                fixed = f'<<Eqn{num}>>'

                            if has_comma:
                                fixed = fixed[:-2] + ',>>'

                            found_placeholders.append((start, end, fixed))

            # Сортируем по позиции и возвращаем только плейсхолдеры
            found_placeholders.sort(key=lambda x: x[0])
            return [placeholder for _, _, placeholder in found_placeholders]

        # Извлекаем из всех частей документа
        for para in document.paragraphs:
            placeholders.extend(extract_from_text(para.text))

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        placeholders.extend(extract_from_text(para.text))

        for section in document.sections:
            for para in section.header.paragraphs:
                placeholders.extend(extract_from_text(para.text))
            for para in section.footer.paragraphs:
                placeholders.extend(extract_from_text(para.text))

        return placeholders

    def process_document(self, doc_path, output_path=None, original_doc_path=None):
        """
        КРИТИЧЕСКИ ИСПРАВЛЕННАЯ обработка документа
        """
        # Сброс статистики
        self.stats = {
            'placeholders_replaced': 0,
            'placeholders_found_in_translation': 0,
            'placeholders_found_in_original': 0,
            'damaged_placeholders_fixed': 0,
            'damaged_placeholders_details': []
        }

        if output_path is None:
            output_path = doc_path.parent / f"{doc_path.stem}_fixed{doc_path.suffix}"

        # Загружаем документы
        try:
            document = Document(doc_path)
        except Exception as e:
            return False, f"Ошибка при открытии документа: {e}"

        if not original_doc_path:
            return False, "Не указан оригинальный документ"

        # Извлекаем плейсхолдеры из оригинала
        original_placeholders = self.extract_placeholders_list(original_doc_path)
        if not original_placeholders:
            return False, f"Не удалось извлечь плейсхолдеры из оригинала"

        self.stats['placeholders_found_in_original'] = len(original_placeholders)
        print(f"\n📊 Найдено плейсхолдеров в оригинале: {len(original_placeholders)}")

        # Счетчик для замены
        placeholder_index = 0

        def replace_placeholder(match):
            """Заменяет найденный плейсхолдер на соответствующий из оригинала"""
            nonlocal placeholder_index

            if placeholder_index < len(original_placeholders):
                replacement = original_placeholders[placeholder_index]
                placeholder_index += 1
                self.stats['placeholders_replaced'] += 1
                return replacement
            else:
                placeholder_index += 1
                return match.group(0)

        def process_text(text, location=""):
            """Обрабатывает текст с заменой плейсхолдеров"""
            if not text:
                return text

            # КРИТИЧЕСКИ ВАЖНО: используем правильный паттерн для замены
            # Находим ВСЕ плейсхолдеры (правильные и поврежденные) и заменяем их

            # Паттерн для поиска всех вариантов плейсхолдеров
            comprehensive_pattern = re.compile(
                r'(?:<<|<\s*<|(?<!\w))Eqn\d+(?:\.eps)?(?:>>|>\s*>|(?=\W)|$)(?:[,>\s]*)?',
                re.IGNORECASE
            )

            result_text = comprehensive_pattern.sub(replace_placeholder, text)
            return result_text

        # Сбрасываем индекс
        placeholder_index = 0

        print("\n🔄 Обработка документа...")

        # Обработка параграфов
        for i, para in enumerate(document.paragraphs):
            if para.text and 'Eqn' in para.text:
                original_text = para.text
                new_text = process_text(para.text, f"Параграф {i+1}")
                if new_text != original_text:
                    para.text = new_text
                    print(f"  ✏️ Параграф {i+1}: обработан")

        # Обработка таблиц
        for t_idx, table in enumerate(document.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        if para.text and 'Eqn' in para.text:
                            original_text = para.text
                            new_text = process_text(para.text, f"Таблица {t_idx+1}")
                            if new_text != original_text:
                                para.text = new_text

        # Обработка колонтитулов
        for s_idx, section in enumerate(document.sections):
            for para in section.header.paragraphs:
                if para.text and 'Eqn' in para.text:
                    para.text = process_text(para.text, f"Заголовок {s_idx+1}")
            for para in section.footer.paragraphs:
                if para.text and 'Eqn' in para.text:
                    para.text = process_text(para.text, f"Подвал {s_idx+1}")

        self.stats['placeholders_found_in_translation'] = placeholder_index

        # Сохранение
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            document.save(output_path)

            message = f"Заменено плейсхолдеров: {self.stats['placeholders_replaced']}/{self.stats['placeholders_found_in_translation']}; Доступно в оригинале: {self.stats['placeholders_found_in_original']}"

            return True, message

        except Exception as e:
            return False, f"Ошибка при сохранении: {e}"

def process_single_file():
    """Обработка одного файла"""

    print("\n" + "="*70)
    print(" EQN RESTORATION BATCH - VERSION V4")
    print(" ОКОНЧАТЕЛЬНАЯ ИСПРАВЛЕННАЯ ВЕРСИЯ")
    print("="*70)

    restorer = PlaceholderRestorer()

    # Запрашиваем переведенный файл
    print("\n1. Путь к ПЕРЕВЕДЕННОМУ файлу:")
    file_path = input(" → ").strip().strip('"\'')
    file_path = Path(file_path)

    if not file_path.exists():
        print(f"\n❌ Файл не найден: {file_path}")
        return

    # Запрашиваем оригинал  
    print("\n2. Путь к ОРИГИНАЛЬНОМУ файлу:")
    original_path = input(" → ").strip().strip('"\'')
    original_path = Path(original_path)

    if not original_path.exists():
        print(f"\n❌ Файл не найден: {original_path}")
        return

    # Обработка
    print("\n🔄 Начинаю обработку...")

    success, message = restorer.process_document(
        file_path,
        original_doc_path=original_path
    )

    if success:
        print(f"\n✅ УСПЕШНО!")
        print(f"📊 {message}")
        output_path = file_path.parent / f"{file_path.stem}_fixed{file_path.suffix}"
        print(f"📄 Результат: {output_path}")
    else:
        print(f"\n❌ ОШИБКА: {message}")

if __name__ == "__main__":
    try:
        process_single_file()
    except KeyboardInterrupt:
        print("\n\nОстановлено пользователем")
    except Exception as e:
        print(f"\n\nКритическая ошибка: {e}")
        import traceback
        traceback.print_exc()
